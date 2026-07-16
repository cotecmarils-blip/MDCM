"""Exportación de informes Word (prioridad: dimensión de costos / OMOC)."""
from __future__ import annotations

from io import BytesIO
import math
from pathlib import Path
from typing import Any, Callable

from django.conf import settings
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from .escenario_service import ESCENARIO_ESTANDAR_NOMBRE, ensure_escenario_estandar
from .evaluacion_service import (
    apply_valores_fallback_escenarios,
    collect_terminal_nodes_for_omoe,
    build_evaluacion_schema,
    load_valores_map,
)
from .informe_math import add_latex_equation, tex_text
from .models import Alternativa, Escenario, NodoArbol, Omoe, Proyecto


def _escenario_informe_costos_for_omoe(omoe: Omoe) -> Escenario:
    """
    Preferir «Estandar» si existe; si no, el primer escenario de la dimensión.
    Solo crea «Estandar» cuando la dimensión aún no tiene ningún escenario
    (evitar ensuciar OMOC oceanográfico con un Estandar vacío).
    """
    esc = Escenario.objects.filter(
        omoe=omoe, nombre__iexact=ESCENARIO_ESTANDAR_NOMBRE,
    ).first()
    if esc:
        return esc
    esc = (
        Escenario.objects.filter(omoe=omoe)
        .order_by('orden', 'id')
        .first()
    )
    if esc:
        return esc
    return ensure_escenario_estandar(omoe)


def _branding_dir() -> Path:
    candidates = [
        Path(settings.BASE_DIR) / 'static' / 'branding',
        Path(settings.BASE_DIR) / 'frontend' / 'public',
        Path(settings.BASE_DIR) / 'frontend' / 'build',
    ]
    for path in candidates:
        if path.is_dir():
            return path
    return candidates[0]


def _find_logo(*names: str) -> Path | None:
    base = _branding_dir()
    roots = [
        base,
        Path(settings.BASE_DIR) / 'static' / 'branding',
        Path(settings.BASE_DIR) / 'frontend' / 'public',
        Path(settings.BASE_DIR) / 'frontend' / 'build',
    ]
    for root in roots:
        for name in names:
            path = root / name
            if path.is_file():
                return path
    return None


_HEADER_LOGO_NAVY = (0x0F, 0x2C, 0x59)


def _logo_stream_for_header(logo_path: Path) -> BytesIO:
    """Convierte cualquier logo a tono navy uniforme (misma paleta en el encabezado)."""
    navy = _HEADER_LOGO_NAVY
    with Image.open(logo_path) as img:
        img = img.convert('RGBA')
        pixels = img.load()
        out = Image.new('RGBA', img.size, (0, 0, 0, 0))
        out_pixels = out.load()
        for y in range(img.height):
            for x in range(img.width):
                r, g, b, a = pixels[x, y]
                if a < 16:
                    continue
                if r > 248 and g > 248 and b > 248:
                    continue
                lum = (0.299 * r + 0.587 * g + 0.114 * b) / 255.0
                if lum > 0.97:
                    continue
                alpha_out = int(a * min(1.0, max(0.45, 1.15 - lum)))
                out_pixels[x, y] = (*navy, alpha_out)
        buf = BytesIO()
        out.save(buf, format='PNG')
        buf.seek(0)
        return buf


DOCUMENT_FONT = 'Arial'
DOCUMENT_FONT_SIZE = 11
TABLE_FONT_SIZE = 9
_REPORT_TABLE_WIDTH_CM = 16.8
_REPORT_PAGE_USABLE_CM = 17.5
_REPORT_PACKING_BUFFER_CM = 0.9
# Si tras una tabla queda menos que esto, la página se considera llena:
# evita arrancar otra tabla en un residuo que en la práctica no alcanza.
_REPORT_MIN_REMAINING_TO_PACK_CM = 7.5
# Si la estimación supera la página por menos que esto, en Word casi siempre
# cabía en UNA página: no inventar un "residuo" en la página siguiente.
_REPORT_NEAR_PAGE_OVERFLOW_CM = 7.5
_CHARS_PER_CM_AT_TABLE_FONT = 5.2


def _set_run_font(run, *, size_pt=DOCUMENT_FONT_SIZE, bold=False, color=None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = DOCUMENT_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), DOCUMENT_FONT)
    rFonts.set(qn('w:hAnsi'), DOCUMENT_FONT)
    rFonts.set(qn('w:cs'), DOCUMENT_FONT)


def _setup_document_fonts(doc: Document) -> None:
    """Fuente base del informe: Arial 11 en estilos estándar."""
    for style_name in ('Normal', 'List Bullet', 'List Number'):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = DOCUMENT_FONT
        style.font.size = Pt(DOCUMENT_FONT_SIZE)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), DOCUMENT_FONT)
        rFonts.set(qn('w:hAnsi'), DOCUMENT_FONT)
        rFonts.set(qn('w:cs'), DOCUMENT_FONT)


def _cell_max_line_len(text: Any) -> int:
    parts = str(text or '—').splitlines() or ['—']
    return max((len(part) for part in parts), default=1)


def _column_content_weights(
    headers: list[str],
    rows: list[list[Any]],
) -> list[float]:
    """Peso por columna según el contenido más largo (cabecera o celdas)."""
    column_count = max(len(headers), 1)
    weights: list[float] = []
    for col_idx in range(column_count):
        samples = [_cell_max_line_len(headers[col_idx] if col_idx < len(headers) else '')]
        for row in rows or []:
            if col_idx < len(row):
                samples.append(_cell_max_line_len(row[col_idx]))
            else:
                samples.append(1)
        # Cabeceras largas no deben forzar columnas enormes si el cuerpo es corto.
        body_max = max(samples[1:], default=1)
        header_len = samples[0]
        effective = max(body_max, min(header_len, max(body_max + 4, 12)))
        # Raíz suave: columnas cortas siguen estrechas; las largas ganan espacio.
        weights.append(max(3.0, math.sqrt(effective) * 2.4))
    return weights


def _compute_column_widths_cm(
    headers: list[str],
    rows: list[list[Any]],
    *,
    total_width_cm: float = _REPORT_TABLE_WIDTH_CM,
) -> list[float]:
    """Anchos fijos proporcionales al contenido, con mínimo y redistribución."""
    weights = _column_content_weights(headers, rows)
    column_count = len(weights)
    if column_count == 0:
        return []
    min_cm = min(1.15, total_width_cm / max(column_count * 1.35, 1))
    max_cm = max(min_cm + 0.4, total_width_cm * 0.42)
    total_weight = sum(weights) or float(column_count)
    widths = [total_width_cm * weight / total_weight for weight in weights]

    # Asegura mínimos quitando ancho de las columnas más generosas.
    for _ in range(column_count * 2):
        deficit = 0.0
        donors: list[int] = []
        for idx, width in enumerate(widths):
            if width < min_cm:
                deficit += min_cm - width
                widths[idx] = min_cm
            elif width > min_cm + 0.25:
                donors.append(idx)
        if deficit <= 1e-9 or not donors:
            break
        donor_extra = sum(widths[i] - min_cm for i in donors)
        if donor_extra <= 1e-9:
            break
        for idx in donors:
            share = (widths[idx] - min_cm) / donor_extra
            widths[idx] = max(min_cm, widths[idx] - deficit * share)

    # Limita máximos y reparte el sobrante.
    overflow = 0.0
    receivers: list[int] = []
    for idx, width in enumerate(widths):
        if width > max_cm:
            overflow += width - max_cm
            widths[idx] = max_cm
        elif width < max_cm - 0.2:
            receivers.append(idx)
    if overflow > 1e-9 and receivers:
        room = sum(max_cm - widths[i] for i in receivers)
        if room > 1e-9:
            for idx in receivers:
                share = (max_cm - widths[idx]) / room
                widths[idx] = min(max_cm, widths[idx] + overflow * share)

    # Normaliza a exactamente el ancho útil de página.
    scale = total_width_cm / max(sum(widths), 1e-9)
    return [round(width * scale, 3) for width in widths]


def _apply_column_widths(table, widths_cm: list[float]) -> None:
    """Fija anchos de columna (desactiva autofit igualitario de Word)."""
    if not widths_cm:
        return
    table.autofit = False
    try:
        table.allow_autofit = False
    except AttributeError:
        pass

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    total_twips = str(int(sum(widths_cm) * 567))  # ~567 twips/cm
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), total_twips)
    tbl_w.set(qn('w:type'), 'dxa')

    # Márgenes internos compactos para ganar espacio útil de texto.
    cell_margin_dxa = '40'  # ~0.07 cm
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_cm):
                break
            width_cm = widths_cm[idx]
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width_cm * 567)))
            tc_w.set(qn('w:type'), 'dxa')
            tc_mar = tc_pr.find(qn('w:tcMar'))
            if tc_mar is None:
                tc_mar = OxmlElement('w:tcMar')
                tc_pr.append(tc_mar)
            for edge in ('top', 'left', 'bottom', 'right'):
                node = tc_mar.find(qn(f'w:{edge}'))
                if node is None:
                    node = OxmlElement(f'w:{edge}')
                    tc_mar.append(node)
                node.set(qn('w:w'), cell_margin_dxa)
                node.set(qn('w:type'), 'dxa')


def _table_matrix_from_docx(table) -> tuple[list[str], list[list[str]]]:
    rows_text = [
        [cell.text for cell in row.cells]
        for row in table.rows
    ]
    if not rows_text:
        return [], []
    return rows_text[0], rows_text[1:]


def _apply_table_borders(
    table,
    *,
    internal_vertical: bool = True,
) -> None:
    """Cuadrícula de tabla con líneas horizontales y verticales."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    def _border(edge: str, *, visible: bool) -> OxmlElement:
        el = OxmlElement(f'w:{edge}')
        if visible:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
        else:
            el.set(qn('w:val'), 'nil')
        return el

    borders = OxmlElement('w:tblBorders')
    for edge in ('left', 'right'):
        borders.append(_border(edge, visible=internal_vertical))
    borders.append(_border('insideV', visible=internal_vertical))
    for edge in ('top', 'bottom', 'insideH'):
        borders.append(_border(edge, visible=True))

    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, size_pt=TABLE_FONT_SIZE, bold=bold)


def _finalize_report_table(
    table,
    *,
    header_bold: bool = True,
    internal_vertical: bool = True,
) -> None:
    headers, body_rows = _table_matrix_from_docx(table)
    _apply_column_widths(table, _compute_column_widths_cm(headers, body_rows))
    _apply_table_borders(table, internal_vertical=internal_vertical)
    rows = table.rows
    # Si la tabla cabe en una página, forzar que Word no la parta a medias.
    # Solo las que exceden una página completa pueden cortarse entre filas.
    keep_table_together = (
        _estimate_table_height_cm(headers, body_rows) <= _REPORT_PAGE_USABLE_CM
    )
    for row_idx, row in enumerate(rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn('w:cantSplit')) is None:
            tr_pr.append(OxmlElement('w:cantSplit'))
        bold = header_bold and row_idx == 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                if keep_table_together:
                    paragraph.paragraph_format.keep_with_next = row_idx < len(rows) - 1
                else:
                    paragraph.paragraph_format.keep_with_next = row_idx == 0
                if not paragraph.runs and paragraph.text:
                    content = paragraph.text
                    paragraph.clear()
                    run = paragraph.add_run(content)
                    _set_run_font(run, size_pt=TABLE_FONT_SIZE, bold=bold)
                else:
                    for run in paragraph.runs:
                        _set_run_font(run, size_pt=TABLE_FONT_SIZE, bold=bold)

    # Espacio de respiro debajo de la tabla para separarla del bloque siguiente.
    # Se evita duplicarlo cuando la tabla se re-finaliza (p. ej. tras combinar celdas).
    nxt = table._element.getnext()
    already_spaced = (
        nxt is not None
        and nxt.tag == qn('w:p')
        and nxt.find(qn('w:r')) is None
    )
    if not already_spaced:
        spacer = OxmlElement('w:p')
        spacer_pr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:line'), '120')
        spacing.set(qn('w:lineRule'), 'exact')
        spacer_pr.append(spacing)
        spacer.append(spacer_pr)
        table._element.addnext(spacer)


def _estimate_table_height_cm(
    headers: list[str],
    rows: list[list[Any]],
    *,
    extra_heading_count: int = 0,
) -> float:
    """Estimación realista con margen leve: empaca si cabe; no parte a medias.

    Debe permitir que dos tablas cortas (p. ej. Cost + Risk, 3 filas) compartan
    página, y a la vez no subestimar tablas densas de ~11 filas.
    """
    body_rows = rows or [['—']]
    widths = _compute_column_widths_cm(headers, body_rows)

    def row_height(values: list[Any], *, header: bool = False) -> float:
        max_lines = 1
        for col_idx, value in enumerate(values):
            width_cm = widths[col_idx] if col_idx < len(widths) else (
                _REPORT_TABLE_WIDTH_CM / max(len(headers), 1)
            )
            chars_per_line = max(7, int(width_cm * _CHARS_PER_CM_AT_TABLE_FONT))
            text = str(value or '—')
            lines = sum(
                max(1, math.ceil(len(part) / chars_per_line))
                for part in text.splitlines() or ['']
            )
            max_lines = max(max_lines, min(lines, 6))
        line_height = 0.42 if header else 0.38
        return max(0.55 if header else 0.50, max_lines * line_height + 0.12)

    table_height = row_height(headers, header=True)
    table_height += sum(row_height(list(row)) for row in body_rows)
    chrome_cm = 1.05 + extra_heading_count * 0.72
    raw = table_height + chrome_cm
    # Tablas densas: margen extra moderado (sin empujarlas artificialmente a >1 página).
    if len(body_rows) >= 8:
        raw += 0.9 + 0.06 * (len(body_rows) - 8)
    return max(raw * 1.08, 0.9 + 0.55 * (1 + len(body_rows)) + chrome_cm)


def _estimate_evaluation_table_height_cm(
    headers: list[str],
    rows: list[list[Any]],
    *,
    extra_heading_count: int = 0,
) -> float:
    """Estimación para matrices de Etapa 2 (sigue sesgada a no partir bloques)."""
    column_count = max(len(headers), 1)
    chars_per_line = max(10, int(88 / column_count))

    def lines_for(values: list[Any]) -> int:
        return max(
            (
                max(1, math.ceil(len(str(value or '—')) / chars_per_line))
                for value in values
            ),
            default=1,
        )

    header_height = max(0.68, min(lines_for(headers), 4) * 0.42 + 0.20)
    body_height = sum(
        max(0.62, min(lines_for(list(row)), 4) * 0.40 + 0.18)
        for row in (rows or [['—']])
    )
    return header_height + body_height + 0.95 + extra_heading_count * 0.65


def _page_used_after_block(
    used_before_cm: float,
    block_height_cm: float,
    *,
    min_remaining_cm: float | None = None,
) -> float:
    """Altura ocupada en la ÚLTIMA página tras colocar un bloque.

    Caso crítico (Tabla 7 ≈1 página, estimación 17.6–19 cm):
    el overflow módulo página dejaba ~0.3–1.5 cm "usados" en una página
    siguiente inventada. Cost creía caber ahí, pero en Word seguía en la
    misma hoja bajo Effectiveness y se partía. Si el exceso sobre una página
    es pequeño, se marca la página como llena.

    `min_remaining_cm` permite empaquetar figuras pequeñas (varias por hoja)
    con un umbral menor que el de las tablas.
    """
    usable = _REPORT_PAGE_USABLE_CM
    if min_remaining_cm is None:
        min_remaining_cm = _REPORT_MIN_REMAINING_TO_PACK_CM
    remaining = max(0.0, usable - used_before_cm)

    if block_height_cm <= remaining + 1e-9:
        used = used_before_cm + block_height_cm
    elif used_before_cm <= 0.05:
        # Arranca al tope: ¿cabe en una página real o solo "casi"?
        overflow = block_height_cm - usable
        if overflow <= _REPORT_NEAR_PAGE_OVERFLOW_CM:
            # En la práctica es una sola página llena (no residuo fantasma).
            return usable
        remainder = overflow % usable
        used = usable if remainder <= 1e-9 else remainder
    else:
        # No debería ocurrir: _begin_table_block ya debió saltar de página.
        used = min(usable, block_height_cm)

    leftover = usable - used
    if 0 < leftover < min_remaining_cm:
        return usable
    return used


def _begin_table_block(
    doc: Document,
    estimated_height_cm: float = 8.0,
    *,
    min_remaining_cm: float | None = None,
    buffer_cm: float | None = None,
) -> None:
    """Regla: título + tabla solo si caben COMPLETOS en el residuo actual.

    1) remaining = página usable − ya usado (contador virtual).
    2) Si estimated + buffer > remaining → page break (salvo página vacía).
    3) Actualizar usado; si el residuo queda < MIN_REMAINING → página llena.

    `min_remaining_cm`/`buffer_cm` se pueden bajar para figuras y así permitir
    3+ gráficos pequeños por hoja (las tablas conservan sus valores mayores).
    """
    usable = _REPORT_PAGE_USABLE_CM
    if buffer_cm is None:
        buffer_cm = _REPORT_PACKING_BUFFER_CM
    if min_remaining_cm is None:
        min_remaining_cm = _REPORT_MIN_REMAINING_TO_PACK_CM
    used_height_cm = getattr(doc, '_report_table_page_used_cm', None)
    if used_height_cm is None:
        used_height_cm = 2.8

    remaining_cm = usable - used_height_cm
    # Residuo útil mínimo: no arrancar en un hueco donde solo entra 1 fila.
    fits_completely = (
        estimated_height_cm + buffer_cm <= remaining_cm
        and remaining_cm >= min_remaining_cm
    )

    if not fits_completely and used_height_cm > 0.05:
        doc.add_page_break()
        used_height_cm = 0.0

    setattr(
        doc,
        '_report_table_page_used_cm',
        _page_used_after_block(
            used_height_cm, estimated_height_cm, min_remaining_cm=min_remaining_cm,
        ),
    )
    setattr(doc, '_report_table_blocks', getattr(doc, '_report_table_blocks', 0) + 1)


def _add_table_caption(
    doc: Document,
    title: str,
    *,
    new_page: bool = True,
    estimated_height_cm: float = 8.0,
) -> None:
    """Título consecutivo; el bloque abre página solo cuando no cabe completo."""
    if new_page:
        _begin_table_block(doc, estimated_height_cm)
    number = getattr(doc, '_report_table_number', 0) + 1
    setattr(doc, '_report_table_number', number)
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    # El título viaja junto a su tabla (no queda huérfano al pie de página).
    fmt.keep_with_next = True
    fmt.keep_together = True
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(4)
    run = paragraph.add_run(f'Tabla {number}. {title}')
    _set_run_font(run, size_pt=9)
    run.italic = True


def _add_header_logos(
    document: Document,
    *,
    subtitle_text: str | None = None,
) -> None:
    """Encabezado compacto: logos navy en una fila y línea horizontal pegada debajo."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement

    section = document.sections[0]
    # Acerca el encabezado al borde superior y deja libre el margen del cuerpo.
    section.header_distance = Cm(0.45)
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        # Word crea un párrafo vacío antes de la tabla del encabezado. Eliminar
        # el nodo completo evita el salto de línea visible antes de los logos.
        element = p._element
        element.getparent().remove(element)

    # ENAP = emblema real de la Escuela Naval (los archivos "Logo_ENAP*.png"
    # planos son sólo la palabra "ENAP" en texto; se usan como último recurso).
    logos = [
        (
            _find_logo(
                'Logo_ENAP_emblem.png',
                'Logo_ENAP_header.png',
                'Logo_ENAP.png',
                'Logo ENAP.png',
            ),
            Cm(1.1),
        ),
        (_find_logo('CotecmarLogo_header.png', 'CotecmarLogo.png'), Cm(0.9)),
        (_find_logo('Logo_CUC_header.png', 'Logo_CUC.png', 'Logo CUC.png'), Cm(0.85)),
    ]
    logos = [(p, h) for p, h in logos if p]
    if logos:
        table = header.add_table(rows=1, cols=len(logos), width=Cm(17))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'nil')
            borders.append(el)
        tbl_pr.append(borders)
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)

        for cell, (logo, logo_h) in zip(table.rows[0].cells, logos):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            try:
                run.add_picture(_logo_stream_for_header(logo), height=logo_h)
            except Exception:
                continue

    _add_header_rule(header)


def _add_header_rule(header) -> None:
    """Línea horizontal navy pegada a los logos (estilo membrete)."""
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(1)
    # El espacio queda después de la línea, no entre la línea y los logos.
    rule.paragraph_format.space_after = Pt(8)
    rule.paragraph_format.line_spacing = Pt(2)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0F2C59')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _omoes_costo(proyecto: Proyecto) -> list[Omoe]:
    return list(
        Omoe.objects.filter(proyecto=proyecto, rama_evaluacion='omoc').order_by(
            'orden', 'nombre_modelo', 'id',
        )
    )


def _calcular_costo_alternativa(
    omoe: Omoe,
    escenario: Escenario,
    valores: dict[str, str],
) -> dict[str, Any]:
    from .simulacion_service import _calcular_dimension

    terminales = collect_terminal_nodes_for_omoe(omoe)
    valores_uso, used_fallback = apply_valores_fallback_escenarios(
        omoe, escenario, terminales, valores,
    )
    valor, detalle = _calcular_dimension(
        omoe, terminales, [escenario], valores_uso, debug_logs=None,
    )
    hojas = []
    for h in (detalle or {}).get('hojas') or []:
        hojas.append({
            'nombre': h.get('nombre'),
            'valor': h.get('utilidad') if 'utilidad' in h else h.get('valor'),
            'peso': h.get('peso'),
        })
    return {
        'omoe_id': omoe.id,
        'omoe_nombre': omoe.nombre_modelo or omoe.codigo or f'Dimensión #{omoe.id}',
        'escenario_id': escenario.id,
        'escenario_nombre': escenario.nombre,
        'valor': round(float(valor or 0), 6),
        'modo_valor_terminal': getattr(omoe, 'modo_valor_terminal', None) or 'valor_bruto',
        'escenario_agregacion': getattr(omoe, 'escenario_agregacion', None) or 'minimo_mejor',
        'hojas': hojas,
        'detalle': detalle,
        'valores_fallback_escenarios': used_fallback,
    }


def build_informe_costos_payload(proyecto: Proyecto) -> dict[str, Any]:
    omoes = _omoes_costo(proyecto)
    if not omoes:
        return {
            'ok': False,
            'detail': (
                'No hay dimensión OMOC (costos) en el proyecto. '
                'Cree una dimensión de costos antes de exportar el informe.'
            ),
            'alternativas': [],
            'dimensiones': [],
        }

    alts = list(Alternativa.objects.filter(proyecto=proyecto).order_by('id'))
    dims_meta = []
    for omoe in omoes:
        esc = _escenario_informe_costos_for_omoe(omoe)
        dims_meta.append({
            'omoe_id': omoe.id,
            'nombre': omoe.nombre_modelo or omoe.codigo,
            'escenario': esc.nombre,
            'modo_valor_terminal': getattr(omoe, 'modo_valor_terminal', None),
            'escenario_agregacion': getattr(omoe, 'escenario_agregacion', None),
        })

    filas = []
    used_any_fallback = False
    for alt in alts:
        valores = load_valores_map(alt.id)
        dims = []
        total = 0.0
        for omoe in omoes:
            esc = _escenario_informe_costos_for_omoe(omoe)
            row = _calcular_costo_alternativa(omoe, esc, valores)
            if row.get('valores_fallback_escenarios'):
                used_any_fallback = True
            dims.append(row)
            total += float(row['valor'] or 0)
        filas.append({
            'alternativa_id': alt.id,
            'nombre': alt.nombre,
            'apodo': alt.apodo or '',
            'dimensiones': dims,
            'total_costo': round(total, 6),
        })

    esc_nombres = [d.get('escenario') for d in dims_meta if d.get('escenario')]
    esc_ref = esc_nombres[0] if len(set(esc_nombres)) == 1 and esc_nombres else 'activo'
    nota = (
        f'Consolidación de costos con el escenario «{esc_ref}» '
        'y valores brutos por ítem de costo.'
    )
    if used_any_fallback:
        nota += (
            ' Algunas celdas del escenario de informe estaban vacías y se rellenaron '
            'temporalmente con valores de otros escenarios de la misma dimensión OMOC.'
        )

    return {
        'ok': True,
        'proyecto': {'id': proyecto.id, 'nombre': proyecto.nombre},
        'nota': nota,
        'dimensiones': dims_meta,
        'alternativas': filas,
        'valores_fallback_escenarios': used_any_fallback,
    }


def build_informe_costos_docx(proyecto: Proyecto) -> bytes:
    payload = build_informe_costos_payload(proyecto)
    doc = Document()
    _setup_document_fonts(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    _add_header_logos(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Informe de costos (OMOC)')
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'Proyecto: {proyecto.nombre}')
    _set_run_font(run, bold=True)

    if not payload.get('ok'):
        p = doc.add_paragraph()
        run = p.add_run(payload.get('detail') or 'Sin datos.')
        _set_run_font(run, color=RGBColor(0xB9, 0x1C, 0x1C))
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    nota = doc.add_paragraph()
    run = nota.add_run(payload['nota'])
    _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))

    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run('1. Resumen por alternativa (escenario Estandar)')
    _set_run_font(run, bold=True)

    dim_names = [d['nombre'] for d in payload['dimensiones']]
    _add_table_caption(doc, 'Resumen de costos por alternativa')
    table = doc.add_table(rows=1, cols=2 + len(dim_names))
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], 'Alternativa', bold=True)
    for i, name in enumerate(dim_names):
        _set_cell_text(hdr[i + 1], name or f'Dim. {i + 1}', bold=True)
    _set_cell_text(hdr[-1], 'Total costo', bold=True)

    for fila in payload['alternativas']:
        row = table.add_row().cells
        _set_cell_text(row[0], fila['nombre'])
        for i, dim in enumerate(fila['dimensiones']):
            _set_cell_text(row[i + 1], f"{dim['valor']:.4f}")
        _set_cell_text(row[-1], f"{fila['total_costo']:.4f}")
    _finalize_report_table(table)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    run = h2.add_run('2. Criterios de cálculo')
    _set_run_font(run, bold=True)

    bullets = [
        'Escenario activo: Estandar.',
        'Valor en nodos terminales: valor bruto (sin transformación u(x)).',
        'Agregación dentro del árbol: suma de valores por dimensión.',
    ]
    for text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        _set_run_font(run)

    if payload['dimensiones']:
        doc.add_paragraph()
        h3 = doc.add_paragraph()
        run = h3.add_run('3. Dimensiones de costo incluidas')
        _set_run_font(run, bold=True)
        for d in payload['dimensiones']:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(
                f"{d['nombre']} · escenario «{d['escenario']}» · "
                f"modo valor={d.get('modo_valor_terminal') or '—'} · "
                f"agregación={d.get('escenario_agregacion') or '—'}"
            )
            _set_run_font(run)

    # Desglose por hojas (valores terminales del escenario Estandar)
    has_hojas = any(
        (dim.get('hojas') or [])
        for fila in payload['alternativas']
        for dim in fila.get('dimensiones') or []
    )
    if has_hojas:
        doc.add_paragraph()
        h4 = doc.add_paragraph()
        run = h4.add_run('4. Desglose por nodos terminales (escenario Estandar)')
        _set_run_font(run, bold=True)
        for fila in payload['alternativas']:
            sub = doc.add_paragraph()
            run = sub.add_run(f"Alternativa: {fila['nombre']}")
            _set_run_font(run, bold=True)
            for dim in fila.get('dimensiones') or []:
                hojas = dim.get('hojas') or []
                if not hojas:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(f"{dim['omoe_nombre']} — total {dim['valor']:.4f}")
                _set_run_font(run, bold=True)
                _add_table_caption(
                    doc,
                    f"Desglose de costos de {fila['nombre']} — {dim['omoe_nombre']}",
                )
                table = doc.add_table(rows=1, cols=2)
                _set_cell_text(table.rows[0].cells[0], 'Nodo / ítem de costo', bold=True)
                _set_cell_text(table.rows[0].cells[1], 'Valor x', bold=True)
                for h in hojas:
                    cells = table.add_row().cells
                    _set_cell_text(cells[0], str(h.get('nombre') or '—'))
                    val = h.get('valor')
                    _set_cell_text(
                        cells[1],
                        f'{float(val):.4f}' if val is not None else '—',
                    )
                _finalize_report_table(table)
                doc.add_paragraph()

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Documento generado por HATD · uso interno Cotecmar / ENAP / Universidad de la Costa')
    _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_heading(
    doc: Document,
    text: str,
    *,
    level: int = 1,
    keep_with_next: bool = False,
) -> None:
    p = doc.add_paragraph()
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    run = p.add_run(text)
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))


def _fmt(value: Any, empty: str = '—') -> str:
    if value is None:
        return empty
    if isinstance(value, float):
        return f'{value:.4f}'
    text = str(value)
    return text if text.strip() else empty


def _fmt_percent(value: Any) -> str:
    if value is None or value == '':
        return '—'
    try:
        n = float(value)
        return f'{n:.2f} %'.replace('.00 %', ' %')
    except (TypeError, ValueError):
        return _fmt(value)


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    title: str | None = None,
    new_page: bool = True,
) -> None:
    estimated_height_cm = _estimate_table_height_cm(headers, rows)
    if title:
        _add_table_caption(
            doc,
            title,
            new_page=new_page,
            estimated_height_cm=estimated_height_cm,
        )
    elif new_page:
        _begin_table_block(doc, estimated_height_cm)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)
    if not rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], 'Sin registros.')
        for i in range(1, len(headers)):
            _set_cell_text(cells[i], '')
        _finalize_report_table(table)
        return
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row[:len(headers)]):
            _set_cell_text(cells[i], _fmt(val))
    _finalize_report_table(table)


def _constantes_display(data: dict | None) -> str:
    data = data or {}
    parts = []
    for key in ('L', 'U', 'k', 'T', 'S', 'M', 'V', 'x0'):
        val = data.get(key)
        if val is not None and str(val).strip() != '':
            parts.append(f'{key}={val}')
    return ', '.join(parts)


def _nodos_dimension(omoe: Omoe) -> list[NodoArbol]:
    return list(
        NodoArbol.objects.filter(omoe=omoe)
        .select_related('tipo_nivel', 'parent')
        .order_by('tipo_nivel__orden', 'parent_id', 'orden_visual', 'id')
    )


def _node_level_orden(nodo: NodoArbol) -> int:
    if nodo.tipo_nivel_id and nodo.tipo_nivel and nodo.tipo_nivel.orden is not None:
        return int(nodo.tipo_nivel.orden)
    return 1


def _node_level_label(nodo: NodoArbol) -> str:
    if nodo.tipo_nivel_id and nodo.tipo_nivel and nodo.tipo_nivel.nombre:
        return nodo.tipo_nivel.nombre.upper()
    return f'NIVEL {_node_level_orden(nodo)}'


def _children_by_parent(nodos: list[NodoArbol]) -> dict[int | None, list[NodoArbol]]:
    grouped: dict[int | None, list[NodoArbol]] = {}
    for nodo in nodos:
        grouped.setdefault(nodo.parent_id, []).append(nodo)
    return grouped


def _concept_paths(nodos: list[NodoArbol]) -> list[list[NodoArbol]]:
    by_parent = _children_by_parent(nodos)
    roots = by_parent.get(None, [])
    paths: list[list[NodoArbol]] = []

    def walk(nodo: NodoArbol, path: list[NodoArbol]) -> None:
        current = [*path, nodo]
        children = by_parent.get(nodo.id, [])
        if not children:
            paths.append(current)
            return
        for child in children:
            walk(child, current)

    for root in roots:
        walk(root, [])
    return paths


def _concept_headers(nodos: list[NodoArbol]) -> list[str]:
    labels_by_orden: dict[int, str] = {}
    for nodo in nodos:
        labels_by_orden.setdefault(_node_level_orden(nodo), _node_level_label(nodo))
    return ['MODELO'] + [
        labels_by_orden[orden]
        for orden in sorted(labels_by_orden)
    ]


def _concept_node_text(
    nodo: NodoArbol,
    config_map: dict[int, Any] | None = None,
    *,
    include_weight: bool = False,
) -> str:
    cfg = (config_map or {}).get(nodo.id)
    aplica = getattr(cfg, 'aplica', nodo.aplica) if cfg else nodo.aplica
    peso = getattr(cfg, 'peso', nodo.peso) if cfg else nodo.peso
    parts = [nodo.nombre or '—']
    if include_weight:
        parts.append(f'w={_fmt_percent(peso)}')
    if aplica is False:
        parts.append('No aplica')
    return '\n'.join(parts)


def _concept_rows(
    omoe: Omoe,
    nodos: list[NodoArbol],
    config_map: dict[int, Any] | None = None,
    *,
    include_weight: bool = False,
) -> list[list[str]]:
    paths = _concept_paths(nodos)
    ordens = sorted({_node_level_orden(n) for n in nodos})
    if not paths:
        return [[omoe.nombre_modelo or omoe.codigo or f'Dimensión #{omoe.id}'] + [''] * len(ordens)]

    rows = []
    for path in paths:
        by_orden = {_node_level_orden(n): n for n in path}
        row = [omoe.nombre_modelo or omoe.codigo or f'Dimensión #{omoe.id}']
        for orden in ordens:
            nodo = by_orden.get(orden)
            row.append(
                _concept_node_text(nodo, config_map, include_weight=include_weight)
                if nodo else ''
            )
        rows.append(row)
    return rows


def _is_estandar_escenario(nombre: str | None) -> bool:
    return (nombre or '').strip().lower() in ('estandar', 'estándar', 'standard')


def _merge_table_column_runs(table, col_idx: int, *, start_row: int = 1) -> None:
    rows = table.rows
    i = start_row
    while i < len(rows):
        j = i + 1
        val = rows[i].cells[col_idx].text.strip()
        while j < len(rows) and rows[j].cells[col_idx].text.strip() == val and val:
            j += 1
        if j - i > 1:
            # Word concatena el contenido de todas las celdas al combinarlas.
            # Vaciar las inferiores evita repetir el mismo texto una vez por fila.
            top_cell = rows[i].cells[col_idx]
            for row_idx in range(i + 1, j):
                rows[row_idx].cells[col_idx].text = ''
            merged = top_cell.merge(rows[j - 1].cells[col_idx])
            _set_cell_text(merged, val)
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        i = j


def _split_path_columns(path: list[NodoArbol]) -> tuple[str, str, NodoArbol]:
    if not path:
        return '', '', None
    if len(path) == 1:
        return '', '', path[0]
    if len(path) == 2:
        return path[0].nombre or '—', '', path[1]
    grupo = path[0].nombre or '—'
    mob = ' / '.join(n.nombre or '—' for n in path[1:-1])
    return grupo, mob, path[-1]


def _terminal_funcion_text(
    terminal: NodoArbol,
    config_map: dict[int, Any] | None = None,
) -> str:
    familia = (
        terminal.familia_funciones
        or _cfg_value(config_map, terminal, 'familia_funciones')
        or ''
    )
    params = _constantes_display(
        terminal.parametros_funcion
        or _cfg_value(config_map, terminal, 'parametros_funcion')
        or {},
    )
    parts = [p for p in (familia, params) if p]
    return ' · '.join(parts) if parts else '—'


def _hierarchy_rows_for_escenario(
    escenario_nombre: str,
    dim_nombre: str,
    nodos: list[NodoArbol],
    config_map: dict[int, Any] | None = None,
) -> list[list[str]]:
    rows = []
    for path in _concept_paths(nodos):
        grupo, mob, terminal = _split_path_columns(path)
        if terminal is None:
            continue
        desc = terminal.descripcion or terminal.observaciones or ''
        funcion = _terminal_funcion_text(terminal, config_map)
        descripcion = desc
        if funcion and funcion != '—':
            descripcion = f'{desc}\nFunción: {funcion}'.strip() if desc else f'Función: {funcion}'
        rows.append([
            escenario_nombre,
            dim_nombre,
            grupo,
            mob,
            terminal.nombre or '—',
            descripcion or '—',
        ])
    return rows


def _peso_rows_for_escenario(
    escenario_nombre: str,
    dim_nombre: str,
    nodos: list[NodoArbol],
    config_map: dict[int, Any] | None = None,
) -> list[list[str]]:
    from .nodo_escenario_service import nodo_effective_aplica, nodo_effective_peso

    active_nodes = [
        nodo for nodo in nodos
        if nodo_effective_aplica(nodo, config_map)
    ]
    active_ids = {nodo.id for nodo in active_nodes}
    by_id = {nodo.id: nodo for nodo in active_nodes}
    children_by_parent: dict[int | None, list[NodoArbol]] = {}
    for nodo in active_nodes:
        children_by_parent.setdefault(nodo.parent_id, []).append(nodo)

    siblings_by_parent: dict[int | None, list[NodoArbol]] = {}
    for nodo in active_nodes:
        siblings_by_parent.setdefault(nodo.parent_id, []).append(nodo)

    def normalized_sibling_fraction(nodo: NodoArbol) -> float:
        siblings = siblings_by_parent.get(nodo.parent_id, [])
        weights = [nodo_effective_peso(sibling, config_map) for sibling in siblings]
        total = sum(weights)
        local = nodo_effective_peso(nodo, config_map)
        if total <= 0:
            return 1.0 / len(siblings) if siblings else 0.0
        return local / total

    accumulated_cache: dict[int, float] = {}

    def accumulated_weight(nodo: NodoArbol) -> float:
        if nodo.id in accumulated_cache:
            return accumulated_cache[nodo.id]
        norm_fraction = normalized_sibling_fraction(nodo)
        if nodo.parent_id in active_ids:
            value = accumulated_weight(by_id[nodo.parent_id]) * norm_fraction
        else:
            value = norm_fraction
        accumulated_cache[nodo.id] = value
        return value

    rows = []
    for nodo in active_nodes:
        parent = by_id.get(nodo.parent_id)
        is_terminal = not children_by_parent.get(nodo.id)
        detail_parts = [
            str(text).strip()
            for text in (nodo.descripcion, nodo.justificacion_peso)
            if str(text or '').strip()
        ]
        if is_terminal:
            funcion = _terminal_funcion_text(nodo, config_map)
            if funcion and funcion != '—':
                detail_parts.append(f'Función de utilidad: {funcion}')
        rows.append([
            escenario_nombre,
            dim_nombre,
            _node_level_label(nodo),
            nodo.nombre or '—',
            parent.nombre if parent else 'Raíz de la dimensión',
            _fmt_percent(nodo_effective_peso(nodo, config_map)),
            _fmt_percent(accumulated_weight(nodo) * 100.0),
            '\n'.join(detail_parts) or '—',
        ])
    return rows


def _add_table_with_merges(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    merge_cols: tuple[int, ...] = (),
    *,
    title: str | None = None,
    new_page: bool = True,
) -> None:
    _add_table(doc, headers, rows, title=title, new_page=new_page)
    if not rows or not merge_cols:
        return
    table = doc.tables[-1]
    for col in merge_cols:
        _merge_table_column_runs(table, col)
    _finalize_report_table(table, internal_vertical=True)


def _cfg_value(config_map: dict[int, Any] | None, nodo: NodoArbol, field: str):
    cfg = (config_map or {}).get(nodo.id)
    if isinstance(cfg, dict):
        return cfg.get(field, getattr(nodo, field, None))
    if cfg is not None:
        return getattr(cfg, field, getattr(nodo, field, None))
    return getattr(nodo, field, None)


def _report_font(size: int, *, bold: bool = False):
    names = (
        'DejaVuSans-Bold.ttf' if bold else 'DejaVuSans.ttf',
        'arialbd.ttf' if bold else 'arial.ttf',
    )
    for name in names:
        try:
            return ImageFont.truetype(name, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _active_concept_nodes(
    nodos: list[NodoArbol],
    config_map: dict[int, Any] | None,
) -> list[NodoArbol]:
    if config_map is None:
        return nodos
    by_id = {n.id: n for n in nodos}
    active_ids = set()
    for nodo in nodos:
        if not bool(_cfg_value(config_map, nodo, 'aplica')):
            continue
        parent = nodo.parent_id
        valid = True
        while parent is not None:
            parent_node = by_id.get(parent)
            if parent_node is None:
                break
            if not bool(_cfg_value(config_map, parent_node, 'aplica')):
                valid = False
                break
            parent = parent_node.parent_id
        if valid:
            active_ids.add(nodo.id)
    return [n for n in nodos if n.id in active_ids]


def _render_concept_map_png(
    omoe: Omoe,
    nodos: list[NodoArbol],
    *,
    config_map: dict[int, Any] | None = None,
    include_weights: bool = False,
) -> BytesIO:
    """Render server-side del mapa conceptual usado en la interfaz."""
    nodes = _active_concept_nodes(nodos, config_map)
    by_parent = _children_by_parent(nodes)
    roots = by_parent.get(None, [])
    level_orders = sorted({_node_level_orden(n) for n in nodes})
    level_index = {order: idx + 1 for idx, order in enumerate(level_orders)}
    max_level = max(level_index.values(), default=1)

    scale = 2
    n_nodes = max(len(nodes), 1)
    # Densidad adaptativa: con muchos nodos el PNG se compacta para no
    # generar figuras de más de una página de alto.
    if n_nodes <= 20:
        density = 1.0
    elif n_nodes <= 50:
        density = 0.82
    elif n_nodes <= 90:
        density = 0.68
    else:
        density = 0.55

    top = int(18 * density) + 4
    margin = int(18 * density) + 6
    node_gap_y = max(3.0, 8.0 * density)
    col_gap = max(36.0, 56.0 * density)
    pad_x = max(6.0, 12.0 * density)
    pad_y = max(4.0, 8.0 * density)
    max_text_w = max(90.0, 130.0 * density)
    label_cap = max(120.0, 180.0 * density)

    font_level = _report_font(max(9, int(13 * density)) * scale, bold=True)
    font_node = _report_font(max(9, int(13 * density)) * scale, bold=True)
    font_root = _report_font(max(11, int(15 * density)) * scale, bold=True)
    font_root_sub = _report_font(max(9, int(12 * density)) * scale, bold=True)

    def _line_h(font) -> float:
        bbox = font.getbbox('Ághjy')
        return (bbox[3] - bbox[1]) / scale + 4

    def _wrap(text: str, font, cap: float) -> list[str]:
        words = str(text or '—').split()
        if not words:
            return ['—']
        lines: list[str] = []
        cur = ''
        for word in words:
            trial = f'{cur} {word}'.strip()
            if not cur or font.getlength(trial) / scale <= cap:
                cur = trial
            else:
                lines.append(cur)
                cur = word
        if cur:
            lines.append(cur)
        return lines

    level_line_h = _line_h(font_level)
    node_line_h = _line_h(font_node)
    root_line_h = _line_h(font_root)

    node_level: dict[int, int] = {
        n.id: level_index.get(_node_level_orden(n), 1) for n in nodes
    }

    def _node_text_lines(n: NodoArbol, cap: float) -> list[str]:
        name = (n.nombre or n.codigo or '—').strip()
        lines = _wrap(name, font_node, cap)
        if include_weights:
            value = _cfg_value(config_map, n, 'peso')
            try:
                num = float(value)
                lines = lines + [f'w={num:.0f}%' if num.is_integer() else f'w={num:.2f}%']
            except (TypeError, ValueError):
                lines = lines + ['w=—']
        return lines

    # 1ª pasada: ancho natural del texto para dimensionar cada columna.
    natural_text_w: dict[int, float] = {}
    for n in nodes:
        lines = _node_text_lines(n, max_text_w)
        natural_text_w[n.id] = max(
            (font_node.getlength(line) / scale for line in lines), default=40
        )

    root_name = omoe.nombre_modelo or omoe.codigo or f'Dimensión #{omoe.id}'
    root_lines = _wrap(root_name, font_root, 180)
    root_text_w = max((font_root.getlength(line) / scale for line in root_lines), default=100)
    root_w = max(150.0, root_text_w + 24)
    root_h = max(48.0, (len(root_lines) + 1) * root_line_h + 12)

    # Etiquetas de nivel (encabezado de cada columna, incl. la dimensión).
    level_label_lines: dict[int, list[str]] = {0: _wrap('DIMENSIÓN', font_level, label_cap)}
    for order in level_orders:
        raw = next(
            (_node_level_label(n) for n in nodes if _node_level_orden(n) == order),
            f'NIVEL {order}',
        )
        level_label_lines[level_index[order]] = _wrap(str(raw).upper(), font_level, label_cap)
    header_lines_max = max((len(v) for v in level_label_lines.values()), default=1)
    header_h = header_lines_max * level_line_h + 12
    content_top = top + header_h + 14

    # Ancho UNIFORME por nivel: se toma la caja más ancha de cada columna.
    col_w: dict[int, float] = {0: root_w}
    for level in range(1, max_level + 1):
        widths = [
            natural_text_w[n.id] + 2 * pad_x for n in nodes
            if node_level[n.id] == level
        ]
        label_w = max(
            (font_level.getlength(line) / scale for line in level_label_lines.get(level, [])),
            default=90,
        )
        col_w[level] = max(max(widths, default=110.0), label_w + 20, 96.0)

    # 2ª pasada: todas las cajas del nivel usan el mismo ancho; el texto se
    # reajusta a ese ancho y se recalcula el alto.
    node_lines: dict[int, list[str]] = {}
    node_w: dict[int, float] = {}
    node_h: dict[int, float] = {}
    for n in nodes:
        level = node_level[n.id]
        box_w = col_w[level]
        lines = _node_text_lines(n, box_w - 2 * pad_x)
        node_lines[n.id] = lines
        node_w[n.id] = box_w
        node_h[n.id] = max(22.0, len(lines) * node_line_h + 2 * pad_y)

    # Posición X (centro) de cada columna, de izquierda a derecha.
    col_center_x: dict[int, float] = {}
    x_cursor = margin
    for level in range(0, max_level + 1):
        col_center_x[level] = x_cursor + col_w[level] / 2
        x_cursor += col_w[level] + col_gap

    # Posición vertical: hojas apiladas (según su alto) y padres centrados.
    y_cursor = content_top
    centers_y: dict[int, float] = {}

    def assign_center_y(nodo: NodoArbol) -> float:
        nonlocal y_cursor
        children = by_parent.get(nodo.id, [])
        if not children:
            cy = y_cursor + node_h[nodo.id] / 2
            y_cursor += node_h[nodo.id] + node_gap_y
        else:
            child_centers = [assign_center_y(child) for child in children]
            cy = sum(child_centers) / len(child_centers)
        centers_y[nodo.id] = cy
        return cy

    for root in roots:
        assign_center_y(root)

    content_bottom = y_cursor
    width = int(x_cursor - col_gap + margin)
    height = int(max(content_bottom, content_top + root_h) + margin)

    image = Image.new('RGB', (width * scale, height * scale), '#ffffff')
    draw = ImageDraw.Draw(image)

    # Raíz de dimensión: centrada verticalmente sobre todo el árbol.
    root_cx = col_center_x[0]
    root_cy = (
        sum(centers_y[r.id] for r in roots) / len(roots)
        if roots else content_top + root_h / 2
    )
    root_x = root_cx - root_w / 2
    root_y = root_cy - root_h / 2

    positions: dict[int, tuple[float, float]] = {}
    for nodo in nodes:
        level = level_index.get(_node_level_orden(nodo), 1)
        positions[nodo.id] = (col_center_x[level], centers_y[nodo.id])

    # Encabezados de nivel arriba de cada columna + separador horizontal.
    draw.line(
        [(margin * scale, (content_top - 10) * scale),
         ((width - margin) * scale, (content_top - 10) * scale)],
        fill='#dce8f4',
        width=1 * scale,
    )
    for level in range(0, max_level + 1):
        cx = col_center_x[level]
        lines = level_label_lines.get(level, [])
        y = top + (header_h - len(lines) * level_line_h) / 2
        for line in lines:
            lw = font_level.getlength(line) / scale
            draw.text(
                ((cx - lw / 2) * scale, y * scale),
                line,
                fill='#6b8caf',
                font=font_level,
            )
            y += level_line_h

    # Conectores tipo peine, ahora en horizontal (izquierda → derecha).
    connector = '#a8c4dc'

    def draw_links(parent_right: float, parent_cy: float, children: list[NodoArbol]) -> None:
        if not children:
            return
        child_points = [
            (positions[ch.id][0] - node_w[ch.id] / 2, positions[ch.id][1])
            for ch in children
        ]
        bus_x = parent_right + col_gap / 2
        draw.line(
            [(parent_right * scale, parent_cy * scale), (bus_x * scale, parent_cy * scale)],
            fill=connector,
            width=2 * scale,
        )
        ys = [y for _, y in child_points]
        draw.line(
            [(bus_x * scale, min(ys + [parent_cy]) * scale),
             (bus_x * scale, max(ys + [parent_cy]) * scale)],
            fill=connector,
            width=2 * scale,
        )
        for child_left, child_cy in child_points:
            draw.line(
                [(bus_x * scale, child_cy * scale), (child_left * scale, child_cy * scale)],
                fill=connector,
                width=2 * scale,
            )
            arrow = [
                (child_left * scale, child_cy * scale),
                ((child_left - 8) * scale, (child_cy - 5) * scale),
                ((child_left - 8) * scale, (child_cy + 5) * scale),
            ]
            draw.polygon(arrow, fill=connector)

    draw_links(root_x + root_w, root_cy, roots)
    for nodo in nodes:
        cx, cy = positions[nodo.id]
        draw_links(cx + node_w[nodo.id] / 2, cy, by_parent.get(nodo.id, []))

    # Raíz de dimensión (azul pastel medio).
    draw.rounded_rectangle(
        [root_x * scale, root_y * scale, (root_x + root_w) * scale, (root_y + root_h) * scale],
        radius=10 * scale,
        fill='#6b9fd4',
        outline='#4f87bc',
        width=2 * scale,
    )
    ry = root_y + 12
    for line in root_lines:
        lw = font_root.getlength(line) / scale
        draw.text(
            ((root_cx - lw / 2) * scale, ry * scale),
            line,
            fill='#ffffff',
            font=font_root,
        )
        ry += root_line_h
    sub = (omoe.rama_evaluacion or 'omoe').upper()
    sub_w = font_root_sub.getlength(sub) / scale
    draw.text(
        ((root_cx - sub_w / 2) * scale, ry * scale),
        sub,
        fill='#e8f2fc',
        font=font_root_sub,
    )

    # Nodos por nivel: azules pastel (más claro hacia las hojas).
    palette = [
        ('#6b9fd4', '#4f87bc', '#ffffff'),
        ('#8eb8e3', '#6b9fd4', '#1e3a5f'),
        ('#b3d0ef', '#8eb8e3', '#1e3a5f'),
        ('#d4e6f7', '#b3d0ef', '#334155'),
        ('#eaf3fb', '#c8dff2', '#475569'),
    ]
    weight_palette = ('#7eb3e0', '#5a96c8', '#1e3a5f')
    for nodo in nodes:
        cx, cy = positions[nodo.id]
        level = level_index.get(_node_level_orden(nodo), 1)
        fill, outline, text_color = palette[min(level - 1, len(palette) - 1)]
        if include_weights:
            fill, outline, text_color = weight_palette
        w = node_w[nodo.id]
        h = node_h[nodo.id]
        x0, y0 = cx - w / 2, cy - h / 2
        draw.rounded_rectangle(
            [x0 * scale, y0 * scale, (x0 + w) * scale, (y0 + h) * scale],
            radius=9 * scale,
            fill=fill,
            outline=outline,
            width=2 * scale,
        )
        lines = node_lines[nodo.id]
        text_block_h = len(lines) * node_line_h
        ty = cy - text_block_h / 2
        for line in lines:
            lw = font_node.getlength(line) / scale
            draw.text(
                ((cx - lw / 2) * scale, ty * scale),
                line,
                fill=text_color,
                font=font_node,
            )
            ty += node_line_h

    output = BytesIO()
    image.save(output, format='PNG', optimize=True)
    output.seek(0)
    return output


def _begin_map_block(doc: Document, estimated_height_cm: float) -> None:
    """Agrupa mapas pequeños en una página y salta solo cuando no caben."""
    usable_height_cm = 21.0
    used_height_cm = getattr(doc, '_report_map_page_used_cm', None)
    if used_height_cm is None:
        # La primera página de mapas ya trae título de sección y texto introductorio.
        used_height_cm = 3.4
    elif used_height_cm + estimated_height_cm > usable_height_cm:
        doc.add_page_break()
        used_height_cm = 0.0

    setattr(
        doc,
        '_report_map_page_used_cm',
        used_height_cm + estimated_height_cm,
    )
    setattr(doc, '_report_map_blocks', getattr(doc, '_report_map_blocks', 0) + 1)


def _concept_map_picture_layout(
    omoe: Omoe,
    nodos: list[NodoArbol],
    *,
    config_map: dict[int, Any] | None = None,
    include_weights: bool = False,
) -> tuple[BytesIO, float, float]:
    picture = _render_concept_map_png(
        omoe,
        nodos,
        config_map=config_map,
        include_weights=include_weights,
    )
    # Bloque en tope de página: título + figura ≤ ~17 cm de alto.
    max_w_cm = 16.0
    max_h_cm = 16.5
    render_scale = 2
    layout_dpi = 96
    with Image.open(picture) as img:
        px_w, px_h = img.size
    picture.seek(0)
    aspect = (px_h / px_w) if px_w else 1.0

    logical_width_px = px_w / render_scale
    natural_width_cm = logical_width_px / layout_dpi * 2.54
    active_node_count = len(_active_concept_nodes(nodos, config_map))
    content_width_cap_cm = min(max_w_cm, 9.0 + min(active_node_count, 35) * 0.2)
    width_cm = min(content_width_cap_cm, natural_width_cm, max_w_cm)
    height_cm = width_cm * aspect
    if height_cm > max_h_cm:
        width_cm = max_h_cm / aspect
        height_cm = max_h_cm

    return picture, width_cm, height_cm


def _add_prepared_concept_map_picture(
    doc: Document,
    picture: BytesIO,
    width_cm: float,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(picture, width=Cm(width_cm))


def _add_concept_map_block(
    doc: Document,
    headings: list[str],
    omoe: Omoe,
    nodos: list[NodoArbol],
    *,
    config_map: dict[int, Any] | None = None,
    include_weights: bool = False,
) -> None:
    picture, width_cm, height_cm = _concept_map_picture_layout(
        omoe,
        nodos,
        config_map=config_map,
        include_weights=include_weights,
    )
    estimated_heading_height_cm = 0.75 * len(headings)
    estimated_block_height_cm = height_cm + estimated_heading_height_cm + 0.7
    _begin_map_block(doc, estimated_block_height_cm)
    for heading in headings:
        _add_heading(doc, heading, level=3, keep_with_next=True)
    _add_prepared_concept_map_picture(doc, picture, width_cm)


def _add_etapa1_arboles_section(
    doc: Document,
    schema: dict[str, Any],
    *,
    progress_callback: Callable[[int, str], None] | None = None,
) -> None:
    """1.1 Árboles dimensionales renderizados como el mapa conceptual."""
    _add_heading(doc, 'Etapa 1. Definición y estructuración del proyecto', level=1)
    _add_heading(doc, '1.1. Gráficas de los árboles dimensionales', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'Los árboles se representan visualmente como en el mapa conceptual del software: '
        'niveles, nodos y conexiones jerárquicas.'
    )
    _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))

    from .nodo_escenario_service import load_config_map

    dimensiones = schema.get('dimensiones') or []
    dimension_ids = [dim['omoe_id'] for dim in dimensiones]
    escenarios_por_dimension: dict[int, list[Escenario]] = {
        dimension_id: [] for dimension_id in dimension_ids
    }
    for escenario in Escenario.objects.filter(
        omoe_id__in=dimension_ids,
    ).order_by('orden', 'nombre', 'id'):
        escenarios_por_dimension.setdefault(escenario.omoe_id, []).append(escenario)
    total_mapas = max(
        sum(
            1 + sum(
                not _is_estandar_escenario(esc.nombre)
                for esc in escenarios_por_dimension.get(dim['omoe_id'], [])
            )
            for dim in dimensiones
        ),
        1,
    )
    mapas_generados = 0

    def map_progress() -> None:
        if progress_callback:
            progress_callback(
                20 + round(35 * mapas_generados / total_mapas),
                f'Generando mapas dimensionales ({mapas_generados}/{total_mapas})',
            )

    for dim in dimensiones:
        omoe = Omoe.objects.get(pk=dim['omoe_id'])
        nodos = _nodos_dimension(omoe)
        dim_label = f"{dim['omoe_nombre']} ({(dim.get('rama_evaluacion') or '').upper() or 'TIPO'})"

        _add_concept_map_block(
            doc,
            [dim_label, 'Árbol estándar / completo (sin pesos)'],
            omoe,
            nodos,
            include_weights=False,
        )
        mapas_generados += 1
        map_progress()

        escenarios = escenarios_por_dimension.get(omoe.id, [])
        for esc in escenarios:
            if _is_estandar_escenario(esc.nombre):
                continue
            config = load_config_map(esc.id)
            _add_concept_map_block(
                doc,
                [f'Escenario: {esc.nombre} (con pesos)'],
                omoe,
                nodos,
                config_map=config,
                include_weights=True,
            )
            mapas_generados += 1
            map_progress()


def _add_alternativas_section(doc: Document, proyecto: Proyecto) -> None:
    alternativas = list(
        Alternativa.objects.filter(proyecto=proyecto)
        .prefetch_related('capacidades', 'caracteristicas__plantilla')
        .order_by('id')
    )
    rows = []
    for alt in alternativas:
        caracts = [
            f'{c.plantilla.nombre}: {c.dato} {c.plantilla.unidad or ""}'.strip()
            for c in alt.caracteristicas.all()
        ]
        caps = [c.nombre for c in alt.capacidades.all()]
        costo = f'{alt.costo} {alt.costo_unidad}' if alt.costo is not None else ''
        rows.append([
            alt.nombre,
            alt.descripcion,
            alt.referencia,
            costo,
            '; '.join(caps),
            '; '.join(caracts),
        ])
    headers = [
        'Alternativa',
        'Descripción',
        'Referencia',
        'Costo declarado',
        'Capacidades',
        'Características',
    ]
    # Título de sección + caption + tabla viajan juntos (sin página huérfana).
    _begin_table_block(
        doc,
        _estimate_table_height_cm(headers, rows, extra_heading_count=1),
    )
    _add_heading(doc, '1.2. Definición de las alternativas', level=2, keep_with_next=True)
    _add_table(
        doc,
        headers,
        rows,
        title='Definición de las alternativas',
        new_page=False,
    )

    for alt in alternativas:
        if not alt.foto:
            continue
        foto_path = Path(alt.foto.path) if alt.foto.name else None
        if not foto_path or not foto_path.is_file():
            continue
        sub = doc.add_paragraph()
        sub.paragraph_format.keep_with_next = True
        run = sub.add_run(f'Imagen — {alt.nombre}')
        _set_run_font(run, bold=True)
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            pic.add_run().add_picture(str(foto_path), height=Cm(4.5))
        except Exception:
            err = doc.add_paragraph()
            run = err.add_run('No se pudo incorporar la imagen de la alternativa.')
            _set_run_font(run, color=RGBColor(0xB9, 0x1C, 0x1C))


def _add_etapa1_estructura_jerarquica(doc: Document, schema: dict[str, Any]) -> None:
    """1.3 Tabla de definición de la estructura jerárquica."""
    from .nodo_escenario_service import load_config_map

    headers = [
        'Escenario',
        'Dimensión',
        'Grupo de afinidad',
        'Nodo intermedio — MOB',
        'Nodo terminal — DT',
        'Descripción o función',
    ]
    section_title_pending = '1.3. Tabla de definición de la estructura jerárquica'
    for dim in schema.get('dimensiones') or []:
        omoe = Omoe.objects.get(pk=dim['omoe_id'])
        nodos = _nodos_dimension(omoe)
        dim_nombre = dim['omoe_nombre'] or omoe.nombre_modelo or omoe.codigo
        escenarios = list(Escenario.objects.filter(omoe=omoe).order_by('orden', 'nombre', 'id'))
        if not escenarios:
            rows = _hierarchy_rows_for_escenario('—', dim_nombre, nodos)
            extra = 2 if section_title_pending else 1
            _begin_table_block(
                doc,
                _estimate_table_height_cm(headers, rows, extra_heading_count=extra),
            )
            if section_title_pending:
                _add_heading(doc, section_title_pending, level=2, keep_with_next=True)
                section_title_pending = None
            _add_heading(doc, dim_nombre, level=3, keep_with_next=True)
            _add_table_with_merges(
                doc,
                headers,
                rows,
                merge_cols=(0, 1, 2, 3),
                title=f'Estructura jerárquica — {dim_nombre}',
                new_page=False,
            )
            continue

        for esc in escenarios:
            config = load_config_map(esc.id)
            rows = _hierarchy_rows_for_escenario(esc.nombre, dim_nombre, nodos, config)
            extra = 2 if section_title_pending else 1
            _begin_table_block(
                doc,
                _estimate_table_height_cm(headers, rows, extra_heading_count=extra),
            )
            if section_title_pending:
                _add_heading(doc, section_title_pending, level=2, keep_with_next=True)
                section_title_pending = None
            _add_heading(
                doc,
                f'{dim_nombre} · Escenario: {esc.nombre}',
                level=3,
                keep_with_next=True,
            )
            _add_table_with_merges(
                doc,
                headers,
                rows,
                merge_cols=(0, 1, 2, 3),
                title=f'Estructura jerárquica — {dim_nombre} — {esc.nombre}',
                new_page=False,
            )
    if section_title_pending:
        _add_heading(doc, section_title_pending, level=2)
        note = doc.add_paragraph()
        run = note.add_run('No hay dimensiones configuradas para mostrar la estructura.')
        _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))


def _add_etapa1_pesos_nodo(doc: Document, schema: dict[str, Any]) -> None:
    """1.4 Tabla de pesos por nodo."""
    from .nodo_escenario_service import load_config_map

    def add_weight_explanation() -> None:
        explanation = doc.add_paragraph()
        explanation.paragraph_format.keep_with_next = True
        _set_run_font(explanation.add_run(
            'El peso local es el valor configurado para el nodo dentro de su grupo de '
            'hermanos. El peso acumulado normaliza esos valores en cada nivel '
            '(peso local / Σ hermanos activos) y los multiplica a lo largo de la ruta '
            'desde la raíz, igual que en el motor de simulación.'
        ))

    headers = [
        'Escenario',
        'Dimensión',
        'Nivel',
        'Nodo',
        'Nodo padre',
        'Peso local',
        'Peso acumulado',
        'Descripción / función',
    ]
    section_title_pending = '1.4. Tabla de pesos por nodo'
    for dim in schema.get('dimensiones') or []:
        omoe = Omoe.objects.get(pk=dim['omoe_id'])
        nodos = _nodos_dimension(omoe)
        dim_nombre = dim['omoe_nombre'] or omoe.nombre_modelo or omoe.codigo
        escenarios = list(Escenario.objects.filter(omoe=omoe).order_by('orden', 'nombre', 'id'))
        if not escenarios:
            rows = _peso_rows_for_escenario('—', dim_nombre, nodos)
            extra = 3 if section_title_pending else 1
            _begin_table_block(
                doc,
                _estimate_table_height_cm(headers, rows, extra_heading_count=extra),
            )
            if section_title_pending:
                _add_heading(doc, section_title_pending, level=2, keep_with_next=True)
                add_weight_explanation()
                section_title_pending = None
            _add_heading(doc, dim_nombre, level=3, keep_with_next=True)
            _add_table_with_merges(
                doc,
                headers,
                rows,
                merge_cols=(0, 1),
                title=f'Pesos por nodo — {dim_nombre}',
                new_page=False,
            )
            continue

        for esc in escenarios:
            config = load_config_map(esc.id)
            rows = _peso_rows_for_escenario(esc.nombre, dim_nombre, nodos, config)
            extra = 3 if section_title_pending else 1
            _begin_table_block(
                doc,
                _estimate_table_height_cm(headers, rows, extra_heading_count=extra),
            )
            if section_title_pending:
                _add_heading(doc, section_title_pending, level=2, keep_with_next=True)
                add_weight_explanation()
                section_title_pending = None
            _add_heading(
                doc,
                f'{dim_nombre} · Escenario: {esc.nombre}',
                level=3,
                keep_with_next=True,
            )
            _add_table_with_merges(
                doc,
                headers,
                rows,
                merge_cols=(0, 1),
                title=f'Pesos por nodo — {dim_nombre} — {esc.nombre}',
                new_page=False,
            )
    if section_title_pending:
        _add_heading(doc, section_title_pending, level=2)
        note = doc.add_paragraph()
        run = note.add_run('No hay dimensiones configuradas para mostrar los pesos.')
        _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))


def _display_valor_evaluacion(raw: Any, col: dict) -> str:
    """Texto legible de la celda de evaluación (valor x de entrada)."""
    if raw is None or str(raw).strip() == '':
        return '—'
    text = str(raw).strip()
    if col.get('input_kind') == 'riesgo' and '|' in text:
        prob, cons = text.split('|', 1)
        if prob and cons:
            return f'P={prob} · C={cons}'
    if col.get('unidad'):
        return f'{text} {col["unidad"]}'.strip()
    return text


def _load_alternativas_valores(proyecto: Proyecto) -> list[tuple[Alternativa, dict[str, str]]]:
    alternativas = list(Alternativa.objects.filter(proyecto=proyecto).order_by('id'))
    return [(alt, load_valores_map(alt.id)) for alt in alternativas]


def _columnas_escenario_dimension(dimension: dict, escenario_id: int) -> list[dict]:
    return [
        col for col in (dimension.get('columnas') or [])
        if col.get('escenario_id') == escenario_id
    ]


def _evaluacion_matrix_headers(
    alternativas_valores: list[tuple[Alternativa, dict[str, str]]],
) -> list[str]:
    return ['Nodo terminal'] + [
        alt.nombre or f'Alternativa #{alt.id}'
        for alt, _ in alternativas_valores
    ]


def _evaluacion_matrix_rows(
    alternativas_valores: list[tuple[Alternativa, dict[str, str]]],
    columnas: list[dict],
) -> list[list[str]]:
    rows: list[list[str]] = []
    for col in columnas:
        row = [col.get('terminal_nombre') or col.get('label') or '—']
        for _, valores in alternativas_valores:
            row.append(_display_valor_evaluacion(valores.get(col['key'], ''), col))
        rows.append(row)
    return rows


def _matriz_bruta_column_label(col: dict, dim_nombre: str) -> str:
    terminal = col.get('terminal_nombre') or col.get('label') or '—'
    escenario = col.get('escenario_nombre') or ''
    if escenario:
        return f'{dim_nombre} — {terminal} ({escenario})'
    return f'{dim_nombre} — {terminal}'


def _scientific_number(value: Any) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return _fmt(value)
    if abs(number) >= 1_000_000 or (number != 0 and abs(number) < 0.0001):
        return f'{number:.6e}'
    return f'{number:.6f}'.rstrip('0').rstrip('.')


def _add_equation(doc: Document, equation: str) -> None:
    """Ecuación centrada como OMML nativo de Word (LaTeX → editable)."""
    add_latex_equation(doc, equation, center=True)


def _walk_calculation_trace(trace: dict[str, Any] | None):
    if not isinstance(trace, dict):
        return
    yield trace
    for child in trace.get('hijos') or []:
        child_trace = child.get('trace') if isinstance(child, dict) else None
        if isinstance(child_trace, dict):
            yield from _walk_calculation_trace(child_trace)
        elif isinstance(child, dict) and child.get('kind'):
            yield from _walk_calculation_trace(child)


def _utility_equation_label(leaf: dict[str, Any]) -> str:
    family = leaf.get('familia_label') or leaf.get('familia') or 'Utilidad'
    utility_type = leaf.get('utilidad_tipo') or ''
    debug = leaf.get('debug') or {}
    direction = 'creciente' if debug.get('is_increasing', True) else 'decreciente'
    if family.startswith('Valor bruto'):
        return 'z(x) = x'
    if family.startswith('Riesgo'):
        return 'r(x) = probabilidad × consecuencia'
    if utility_type == 'LinearUtilityFunction':
        base = 'n(x) = clip((x−L)/(U−L), 0, 1)'
        return f'{base}; u(x) = {"n(x)" if direction == "creciente" else "1−n(x)"}'
    if utility_type == 'ExponentialUtilityFunction':
        return 'u(x) = (exp(k·n(x))−1)/(exp(k)−1)'
    if utility_type == 'LogarithmicUtilityFunction':
        return 'u(x) = ln(1+k·n(x))/ln(1+k)'
    if utility_type == 'SigmoidalUtilityFunction':
        return 'u(x) = 1/(1+exp(−k·(n(x)−m)))'
    if utility_type == 'DiscreteUtilityFunction':
        return 'u(x) = valor de la escala discreta configurada'
    return f'u(x) — {family}'


def _leaf_value_substitution(
    leaf: dict[str, Any],
    scenario: dict[str, Any],
) -> str:
    x = scenario.get('x')
    result = _scientific_number(scenario.get('u'))
    family = leaf.get('familia_label') or ''
    if family.startswith('Valor bruto'):
        return f'z({_fmt(x)}) = {_fmt(x)}'
    if family.startswith('Riesgo'):
        parts = str(x or '').split('|', 1)
        if len(parts) == 2:
            return f'r = {_fmt(parts[0])} × {_fmt(parts[1])} = {result}'
        return f'r({_fmt(x)}) = {result}'
    debug = leaf.get('debug') or {}
    utility_type = leaf.get('utilidad_tipo') or ''
    if utility_type == 'LinearUtilityFunction':
        try:
            x_num = float(x)
            lower = float(debug.get('L'))
            upper = float(debug.get('U'))
            normalized = (
                max(0.0, min(1.0, (x_num - lower) / (upper - lower)))
                if upper != lower else 0.0
            )
            base = normalized if debug.get('is_increasing', True) else 1.0 - normalized
            return (
                f'n=clip(({_scientific_number(x_num)}−{_scientific_number(lower)})/'
                f'({_scientific_number(upper)}−{_scientific_number(lower)}))'
                f'={_scientific_number(normalized)}; '
                f'u={_scientific_number(base)}'
            )
        except (TypeError, ValueError):
            pass
    return (
        f'{_utility_equation_label(leaf)}; '
        f'x={_fmt(x)} ⇒ resultado={result}'
    )


def _rollup_substitution(trace: dict[str, Any]) -> str:
    children = trace.get('hijos') or []
    if not children:
        return '—'
    def child_value(child: dict[str, Any]):
        return child.get('valor') if child.get('valor') is not None else child.get('utilidad')

    values = [_scientific_number(child_value(child)) for child in children]
    if str(trace.get('formula') or '').startswith('Σ(x'):
        return ' + '.join(values) + f' = {_scientific_number(trace.get("valor"))}'
    numerator = ' + '.join(
        f'({_scientific_number(child_value(child))}'
        f' × {_scientific_number(child.get("peso"))})'
        for child in children
    )
    denominator = _scientific_number(
        trace.get('suma_pesos')
        if trace.get('suma_pesos') is not None
        else sum(float(child.get('peso') or 0) for child in children)
    )
    return (
        f'[{numerator}] / {denominator}'
        f' = {_scientific_number(trace.get("valor"))}'
    )


def _pick_worked_example(
    resultados: list[dict[str, Any]],
) -> tuple[dict[str, Any], dict[str, Any]] | None:
    """Prefiere una dimensión en modo utilidad con traza y varios terminales."""
    candidates: list[tuple[int, dict[str, Any], dict[str, Any]]] = []
    for alt in resultados:
        for dim in alt.get('dimensiones') or []:
            trace = ((dim.get('detalle') or {}).get('trace')) or {}
            if not trace:
                continue
            leaves = [
                node for node in _walk_calculation_trace(trace)
                if node.get('kind') == 'leaf'
            ]
            modo = str(trace.get('modo_valor_terminal') or '')
            score = len(leaves) * 10 + (5 if modo == 'utilidad' else 0)
            candidates.append((score, alt, dim))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    _, alt, dim = candidates[0]
    return alt, dim


def _add_body_paragraph(
    doc: Document,
    text: str,
    *,
    indent: bool = True,
    keep_with_next: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.5)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True
    _set_run_font(paragraph.add_run(text))


def _add_step_heading(doc: Document, text: str) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(text)
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))


def _add_substep_heading(doc: Document, text: str) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(2)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(text)
    _set_run_font(run, bold=True)


def _add_numbered_equation(doc: Document, equation: str, number: int) -> None:
    """Ecuación numerada OMML nativa (LaTeX → editable en Word)."""
    add_latex_equation(doc, equation, number=number, center=True)


def _leaf_is_quantitative(leaf: dict[str, Any]) -> bool:
    utility_type = leaf.get('utilidad_tipo') or ''
    debug = leaf.get('debug') or {}
    if utility_type == 'DiscreteUtilityFunction':
        return False
    if utility_type == 'LinearUtilityFunction' and debug.get('L') is not None and debug.get('U') is not None:
        return True
    family = str(leaf.get('familia_label') or '')
    if family.startswith('Riesgo') or family.startswith('Valor bruto'):
        return False
    return utility_type in {
        'LinearUtilityFunction',
        'ExponentialUtilityFunction',
        'LogarithmicUtilityFunction',
        'SigmoidalUtilityFunction',
    }


def _leaf_interval_text(leaf: dict[str, Any]) -> str:
    debug = leaf.get('debug') or {}
    if debug.get('L') is not None and debug.get('U') is not None:
        return f'[{_scientific_number(debug["L"])}, {_scientific_number(debug["U"])}]'
    return '—'


def _scenario_xu_text(leaf: dict[str, Any], scenario: dict[str, Any]) -> str:
    x = scenario.get('x')
    u = _scientific_number(scenario.get('u'))
    if x is None or str(x).strip() == '':
        return f'— → {u}'
    text = str(x).strip()
    if (leaf.get('familia_label') or '').startswith('Riesgo') and '|' in text:
        prob, cons = text.split('|', 1)
        return f'P={prob}·C={cons} → {u}'
    try:
        float(text)
        return f'{_scientific_number(text)} → {u}'
    except (TypeError, ValueError):
        return f'{text} → {u}'


def _scenario_names_ordered(leaves: list[dict[str, Any]]) -> list[str]:
    names: list[str] = []
    seen: set[str] = set()
    for leaf in leaves:
        for scenario in leaf.get('escenarios') or []:
            name = scenario.get('nombre') or '—'
            if name not in seen:
                seen.add(name)
                names.append(name)
    return names


def _collect_rollup_nodes(trace: dict[str, Any]) -> list[dict[str, Any]]:
    """Nodos compuestos en orden bottom-up (raíz al final) para el relato."""
    ranked: list[tuple[int, dict[str, Any]]] = []

    def visit(node: dict[str, Any], depth: int) -> None:
        if not isinstance(node, dict):
            return
        for child in node.get('hijos') or []:
            child_trace = child.get('trace') if isinstance(child, dict) else None
            if isinstance(child_trace, dict):
                visit(child_trace, depth + 1)
            elif isinstance(child, dict) and child.get('kind'):
                visit(child, depth + 1)
        if node.get('kind') in ('rollup', 'dimension') and node.get('hijos'):
            ranked.append((depth, node))

    visit(trace, 0)
    ranked.sort(key=lambda item: item[0], reverse=True)
    return [node for _, node in ranked]


def _rollup_children_names(node: dict[str, Any]) -> str:
    names = [child.get('nombre') or '—' for child in (node.get('hijos') or [])]
    if not names:
        return '—'
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f'{names[0]} y {names[1]}'
    return ', '.join(names[:-1]) + f' y {names[-1]}'


def _rollup_weights_text(node: dict[str, Any]) -> str:
    weights = [
        f'{_scientific_number(child.get("peso"))} %'
        for child in (node.get('hijos') or [])
    ]
    return ', '.join(weights) if weights else '—'


def _worked_linear_example(
    leaf: dict[str, Any],
    scenario: dict[str, Any],
) -> tuple[str, str, str] | None:
    debug = leaf.get('debug') or {}
    if leaf.get('utilidad_tipo') != 'LinearUtilityFunction':
        return None
    if debug.get('L') is None or debug.get('U') is None:
        return None
    try:
        x_num = float(scenario.get('x'))
        lower = float(debug['L'])
        upper = float(debug['U'])
    except (TypeError, ValueError):
        return None
    normalized = (
        max(0.0, min(1.0, (x_num - lower) / (upper - lower)))
        if upper != lower else 0.0
    )
    result = normalized if debug.get('is_increasing', True) else 1.0 - normalized
    leaf_name = leaf.get('nombre') or 'j'
    esc_name = scenario.get('nombre') or 'm'
    eq_n = (
        f'n_{{{tex_text(leaf_name)},{tex_text(esc_name)}}} = '
        f'\\mathrm{{clip}}\\left('
        f'\\dfrac{{{_scientific_number(x_num)}-{_scientific_number(lower)}}}'
        f'{{{_scientific_number(upper)}-{_scientific_number(lower)}}}'
        f',\\,0,\\,1\\right) = {_scientific_number(normalized)}'
    )
    eq_u = (
        f'u_{{{tex_text(leaf_name)},{tex_text(esc_name)}}} = '
        f'{_scientific_number(result)}'
    )
    return leaf_name, eq_n, eq_u


def _worked_discrete_example(
    leaf: dict[str, Any],
    scenario: dict[str, Any],
) -> tuple[str, str] | None:
    if leaf.get('utilidad_tipo') != 'DiscreteUtilityFunction' and _leaf_is_quantitative(leaf):
        return None
    x = scenario.get('x')
    if x is None or str(x).strip() == '':
        return None
    try:
        float(x)
        return None
    except (TypeError, ValueError):
        pass
    leaf_name = leaf.get('nombre') or 'j'
    esc_name = scenario.get('nombre') or 'm'
    eq = (
        f'x_{{{tex_text(leaf_name)},{tex_text(esc_name)}}} = '
        f'{tex_text(str(x).strip())}'
        f'\\ \\Rightarrow\\ '
        f'u_{{{tex_text(leaf_name)},{tex_text(esc_name)}}} = '
        f'{_scientific_number(scenario.get("u"))}'
    )
    return leaf_name, eq


def _add_methodology_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    title: str,
) -> None:
    if not rows:
        return
    _begin_table_block(
        doc,
        _estimate_evaluation_table_height_cm(headers, rows, extra_heading_count=0),
    )
    _add_table(doc, headers, rows, title=title, new_page=False)


def _add_matriz_bruta_scientific_explanation(
    doc: Document,
    resultados: list[dict[str, Any]],
    *,
    progress_callback: Callable[[int, str], None] | None = None,
) -> None:
    """Explicación estilo artículo científico: formulación + un ejemplo resuelto."""
    if progress_callback:
        progress_callback(94, 'Escribiendo metodología de la matriz general')

    n_alts = len(resultados)
    n_dims = max((len(alt.get('dimensiones') or []) for alt in resultados), default=0)
    eq = 0

    def next_eq(equation: str) -> None:
        nonlocal eq
        eq += 1
        _add_numbered_equation(doc, equation, eq)

    _add_heading(
        doc,
        '2.3. Metodología del cálculo de la matriz general',
        level=2,
        keep_with_next=True,
    )
    _add_body_paragraph(
        doc,
        'La matriz general X es una única matriz de orden '
        f'{n_alts} × {n_dims} (alternativas × dimensiones). Cada celda xₐd '
        'representa el valor agregado de la dimensión d para la alternativa a, '
        'antes de Pareto, normalización o MADM. A continuación se formaliza el '
        'procedimiento y se resuelve un único ejemplo numérico representativo.',
    )

    example = _pick_worked_example(resultados)
    if example is None:
        _add_heading(doc, '2.3.1. Formulación general', level=3, keep_with_next=True)
        next_eq(r'u_{jm}=f_j(x_{jm})')
        next_eq(r'\bar{u}_j=\dfrac{\sum_m (p_m\cdot u_{jm})}{\sum_m p_m}')
        next_eq(r'U_k=\dfrac{\sum_j (w_{jk}\cdot U_j)}{\sum_j w_{jk}}')
        next_eq(r'X=[x_{ad}]_{a=1,\ldots,A,\ d=1,\ldots,D}')
        if progress_callback:
            progress_callback(97, 'Metodología documentada')
        return

    alt, dim = example
    alt_name = alt.get('nombre') or f"Alternativa #{alt.get('id')}"
    dim_name = dim.get('omoe_nombre') or f"Dimensión #{dim.get('omoe_id')}"
    detail = dim.get('detalle') or {}
    trace = detail.get('trace') or {}
    final_value = _scientific_number(dim.get('valor'))
    modo = trace.get('modo_valor_terminal') or 'utilidad'
    agregacion = dim.get('escenario_agregacion') or 'compensatorio'
    leaves = [
        node for node in _walk_calculation_trace(trace)
        if node.get('kind') == 'leaf'
    ]
    scenario_names = _scenario_names_ordered(leaves)
    n_scenarios = len(scenario_names)
    n_leaves = len(leaves)
    quant_leaves = [leaf for leaf in leaves if _leaf_is_quantitative(leaf)]
    discrete_leaves = [leaf for leaf in leaves if not _leaf_is_quantitative(leaf)]
    modo_label = (
        'modo utilidad' if modo == 'utilidad'
        else 'modo valor bruto' if modo == 'valor_bruto'
        else str(modo)
    )
    agregacion_label = (
        'compensatoria' if agregacion == 'compensatorio'
        else str(agregacion).replace('_', ' ')
    )

    _add_heading(
        doc,
        f'2.3.2. Ejemplo numérico: cálculo de {dim_name} para la alternativa {alt_name}',
        level=3,
        keep_with_next=True,
    )
    _add_body_paragraph(
        doc,
        'Con el propósito de ilustrar el procedimiento utilizado para construir '
        'la matriz general de evaluación, se desarrolla el cálculo de la celda '
        f'correspondiente a la dimensión {dim_name} de la alternativa {alt_name}:',
    )
    next_eq(f'x_{{{tex_text(alt_name)},{tex_text(dim_name)}}}')

    _add_body_paragraph(
        doc,
        f'Para esta evaluación, los nodos terminales operan en {modo_label}, '
        f'la agregación entre escenarios es {agregacion_label} y la propagación '
        'de los resultados dentro del árbol de decisión se realiza mediante un '
        'modelo jerárquico de valor multiatributo (Multi-Attribute Value Theory, '
        'MAVT).',
    )
    _add_body_paragraph(
        doc,
        'El procedimiento comprende tres etapas: transformación de los nodos '
        'terminales, agregación jerárquica y asignación del resultado en la '
        'matriz general.',
    )

    # —— Paso 1 ——
    _add_step_heading(doc, 'Paso 1. Transformación de los nodos terminales')
    _add_body_paragraph(
        doc,
        f'Para cada criterio terminal (j) y escenario operacional (m), el valor '
        f'original x_jm se transforma en una utilidad adimensional u_jm, definida '
        f'en el intervalo [0, 1]. En este ejemplo intervienen {n_leaves} nodos '
        f'terminales'
        + (f' y {n_scenarios} escenarios.' if n_scenarios else '.'),
    )

    if quant_leaves:
        _add_body_paragraph(
            doc,
            'En los criterios cuantitativos se utiliza, cuando corresponde, una '
            'función de normalización lineal:',
        )
        next_eq(
            r'n_{jm}=\mathrm{clip}\left(\dfrac{x_{jm}-L_j}{U_j-L_j},\,0,\,1\right)'
        )
        _add_body_paragraph(
            doc,
            'donde L_j y U_j representan, respectivamente, los límites inferior y '
            'superior definidos para el criterio j. La función clip(·) restringe '
            'el resultado al intervalo [0, 1]:',
        )
        next_eq(r'\mathrm{clip}(z,0,1)=\min\{1,\max\{0,z\}\}')
        _add_body_paragraph(
            doc,
            'Si el criterio es de beneficio (creciente), la utilidad coincide con '
            'el valor normalizado; si es de costo (decreciente), u_jm = 1 − n_jm:',
        )
        next_eq(
            r'u_{jm}=n_{jm}\ \mathrm{(beneficio)}\quad\mathrm{o}\quad'
            r'u_{jm}=1-n_{jm}\ \mathrm{(costo)}'
        )

        worked = None
        worked_esc = None
        for leaf in quant_leaves:
            for scenario in leaf.get('escenarios') or []:
                candidate = _worked_linear_example(leaf, scenario)
                if candidate:
                    worked = candidate
                    worked_esc = scenario.get('nombre') or 'm'
                    break
            if worked:
                break
        if worked:
            leaf_name, eq_n, eq_u = worked
            _add_body_paragraph(
                doc,
                f'Por ejemplo, para el criterio {leaf_name} de la alternativa '
                f'{alt_name} en el escenario {worked_esc}, se obtiene:',
            )
            next_eq(eq_n)
            _add_body_paragraph(doc, 'Por tanto:')
            next_eq(eq_u)

    if discrete_leaves:
        _add_body_paragraph(
            doc,
            'En los criterios cualitativos o discretos, la utilidad se obtiene '
            'directamente de la escala configurada para cada categoría:',
        )
        next_eq(r'u_{jm}=s_j(x_{jm})')
        _add_body_paragraph(
            doc,
            'donde s_j(·) relaciona cada categoría cualitativa con su valor de '
            'utilidad correspondiente.',
        )
        disc_example = None
        for leaf in discrete_leaves:
            for scenario in leaf.get('escenarios') or []:
                disc_example = _worked_discrete_example(leaf, scenario)
                if disc_example:
                    break
            if disc_example:
                break
        if disc_example:
            leaf_name, eq_disc = disc_example
            _add_body_paragraph(
                doc,
                f'Por ejemplo, para el criterio {leaf_name}:',
            )
            next_eq(eq_disc)

    if n_scenarios > 1 and agregacion == 'compensatorio':
        _add_body_paragraph(
            doc,
            f'Una vez transformados los valores de los {n_scenarios} escenarios, '
            'la utilidad agregada del criterio terminal j se calcula mediante:',
        )
        next_eq(
            rf'\bar{{u}}_j='
            rf'\dfrac{{\sum_{{m=1}}^{{{n_scenarios}}} (p_m\cdot u_{{jm}})}}'
            rf'{{\sum_{{m=1}}^{{{n_scenarios}}} p_m}}'
        )
        _add_body_paragraph(
            doc,
            'donde p_m corresponde al peso asignado al escenario m. Si los '
            'escenarios poseen la misma ponderación, la expresión equivale al '
            'promedio aritmético de las utilidades.',
        )

    # Compact transformation tables
    if quant_leaves and scenario_names:
        headers = ['Criterio', 'Intervalo'] + [
            f'{name}: x → u' for name in scenario_names
        ] + ['Utilidad agregada']
        rows = []
        for leaf in quant_leaves:
            by_name = {
                (scenario.get('nombre') or '—'): scenario
                for scenario in (leaf.get('escenarios') or [])
            }
            row = [leaf.get('nombre') or '—', _leaf_interval_text(leaf)]
            for name in scenario_names:
                scenario = by_name.get(name)
                row.append(_scenario_xu_text(leaf, scenario) if scenario else '—')
            row.append(_scientific_number(leaf.get('valor')))
            rows.append(row)
        _add_methodology_table(
            doc,
            headers,
            rows,
            title='Transformación de los criterios terminales cuantitativos',
        )

    if discrete_leaves and scenario_names:
        headers = ['Criterio'] + [
            f'{name}: categoría → utilidad' for name in scenario_names
        ] + ['Utilidad agregada']
        rows = []
        for leaf in discrete_leaves:
            by_name = {
                (scenario.get('nombre') or '—'): scenario
                for scenario in (leaf.get('escenarios') or [])
            }
            row = [leaf.get('nombre') or '—']
            for name in scenario_names:
                scenario = by_name.get(name)
                row.append(_scenario_xu_text(leaf, scenario) if scenario else '—')
            row.append(_scientific_number(leaf.get('valor')))
            rows.append(row)
        _add_methodology_table(
            doc,
            headers,
            rows,
            title='Transformación de los criterios terminales discretos',
        )

    if scenario_names:
        _add_body_paragraph(
            doc,
            'Los escenarios operacionales considerados son:',
        )
        for index, name in enumerate(scenario_names, start=1):
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.paragraph_format.space_after = Pt(2)
            _set_run_font(bullet.add_run(f'{name}'))

    if leaves:
        vector = ',\\,'.join(_scientific_number(leaf.get('valor')) for leaf in leaves)
        _add_body_paragraph(
            doc,
            f'Como resultado de esta primera etapa, se obtiene el vector de '
            f'utilidades agregadas de los {n_leaves} nodos terminales:',
        )
        next_eq(
            rf'\bar{{u}}_{{{tex_text(alt_name)}}}='
            rf'\left[{vector}\right]'
        )

    # —— Paso 2 ——
    rollup_nodes = _collect_rollup_nodes(trace)
    _add_step_heading(doc, 'Paso 2. Agregación jerárquica')
    _add_body_paragraph(
        doc,
        f'Las utilidades agregadas de los nodos terminales se propagan desde '
        f'los niveles inferiores del árbol de valor hasta la raíz de la '
        f'dimensión {dim_name}. Para cada nodo compuesto k se utiliza una suma '
        f'ponderada normalizada:',
    )
    next_eq(r'U_k=\dfrac{\sum_j (w_{jk}\cdot U_j)}{\sum_j w_{jk}}')
    _add_body_paragraph(
        doc,
        'donde U_j es la utilidad del nodo hijo j, w_jk es su peso local y la '
        'suma se extiende sobre los hijos activos del nodo compuesto k.',
    )

    letters = 'abcdefghijklmnopqrstuvwxyz'
    for index, node in enumerate(rollup_nodes):
        node_name = node.get('nombre') or f'Nodo {index + 1}'
        letter = letters[index] if index < len(letters) else str(index + 1)
        children = node.get('hijos') or []
        weights_txt = _rollup_weights_text(node)
        _add_substep_heading(doc, f'{letter}) Nodo {node_name}')
        _add_body_paragraph(
            doc,
            f'El nodo {node_name} integra {_rollup_children_names(node)}, '
            f'con ponderaciones {weights_txt}:',
        )
        # Detailed substitution
        if children:
            terms = ' + '.join(
                f'{_scientific_number(child.get("valor"))}'
                f'\\cdot {_scientific_number(child.get("peso"))}'
                for child in children
            )
            denom = _scientific_number(
                node.get('suma_pesos')
                if node.get('suma_pesos') is not None
                else sum(float(child.get('peso') or 0) for child in children)
            )
            next_eq(
                f'U_{{{tex_text(node_name)}}}='
                f'\\dfrac{{{terms}}}{{{denom}}}='
                f'{_scientific_number(node.get("valor"))}'
            )

    if rollup_nodes:
        headers = [
            'Nodo compuesto',
            'Nodos integrados',
            'Ponderaciones',
            'Utilidad resultante',
        ]
        rows = [
            [
                node.get('nombre') or '—',
                _rollup_children_names(node),
                _rollup_weights_text(node),
                _scientific_number(node.get('valor')),
            ]
            for node in rollup_nodes
        ]
        _add_methodology_table(
            doc,
            headers,
            rows,
            title=f'Agregación jerárquica de la dimensión {dim_name}',
        )

    # —— Paso 3 ——
    _add_step_heading(doc, 'Paso 3. Asignación del resultado en la matriz general')
    _add_body_paragraph(
        doc,
        'El valor obtenido en la raíz del árbol se asigna directamente a la '
        'celda correspondiente de la matriz general X:',
    )
    next_eq(
        f'x_{{{tex_text(alt_name)},{tex_text(dim_name)}}}={final_value}'
    )
    _add_body_paragraph(
        doc,
        'Este procedimiento se aplica de manera análoga a cada combinación de '
        'alternativa y dimensión:',
    )
    next_eq(r'x_{ad},\quad a=1,\ldots,A,\quad d=1,\ldots,D')
    _add_body_paragraph(
        doc,
        'De esta forma, la repetición del proceso permite construir la matriz '
        'general:',
    )
    next_eq(
        r'X=\begin{bmatrix}'
        r'x_{11} & x_{12} & \cdots & x_{1D} \\'
        r'x_{21} & x_{22} & \cdots & x_{2D} \\'
        r'\vdots & \vdots & \ddots & \vdots \\'
        r'x_{A1} & x_{A2} & \cdots & x_{AD}'
        r'\end{bmatrix}'
    )
    _add_body_paragraph(
        doc,
        'En la sección siguiente se presenta únicamente la matriz general '
        'resultante. Los cálculos correspondientes a las demás celdas no se '
        'reproducen individualmente, debido a que siguen la misma estructura '
        f'de transformación, agregación de escenarios y propagación jerárquica '
        f'desarrollada para el par ({alt_name}, {dim_name}).',
    )

    if progress_callback:
        progress_callback(97, 'Metodología documentada')


def _build_matriz_bruta_calculo(
    proyecto: Proyecto,
    progress_callback: Callable[[int, str], None] | None = None,
    details_out: list[dict[str, Any]] | None = None,
) -> tuple[list[str], list[list[Any]], str | None]:
    """
    Matriz bruta del cálculo (antes de Pareto / normalización / MADM):
    filas = alternativas, columnas = dimensiones, celdas = valor agregado.
    """
    from django.core.exceptions import ValidationError

    from .madm_pipeline import build_matrix_from_rollups
    from .simulacion_service import _rollup_alternativas_simulacion, validar_simulacion

    if progress_callback:
        progress_callback(90, 'Validando datos para la matriz general')
    validacion = validar_simulacion(proyecto)
    if not validacion.get('ok'):
        total = validacion.get('total_faltantes') or 0
        return [], [], (
            f'No se pudo construir la matriz bruta: faltan {total} valor(es) '
            f'de evaluación. Complete la matriz de entradas y vuelva a exportar.'
        )

    try:
        if progress_callback:
            progress_callback(91, 'Calculando valores agregados por dimensión')
        resultados, dimensiones_meta = _rollup_alternativas_simulacion(proyecto)
        if progress_callback:
            progress_callback(93, 'Construyendo la matriz general')
        matrix, alt_names, dim_names, directions = build_matrix_from_rollups(
            resultados,
            dimensiones_meta,
        )
        if details_out is not None:
            details_out.extend(resultados)
    except ValidationError as exc:
        detail = exc.messages[0] if getattr(exc, 'messages', None) else str(exc)
        return [], [], f'No se pudo construir la matriz bruta: {detail}'
    except Exception as exc:
        return [], [], f'No se pudo construir la matriz bruta: {exc}'

    headers = ['Alternativa'] + list(dim_names)
    # Fila de sentido de preferencia (referencia) + filas de valores.
    pref_row = ['Sentido'] + [
        ('Maximizar' if str(d).lower() in ('max', 'benefit', 'beneficio') else 'Minimizar')
        for d in directions
    ]
    rows: list[list[Any]] = [pref_row]
    for i, name in enumerate(alt_names):
        rows.append([name] + [round(float(v), 6) for v in matrix[i]])
    if progress_callback:
        progress_callback(93, 'Matriz general calculada')
    return headers, rows, None


def _reset_report_table_packing(doc: Document) -> None:
    setattr(doc, '_report_table_page_used_cm', None)


def _add_etapa2_evaluaciones(
    doc: Document,
    proyecto: Proyecto,
    schema: dict[str, Any],
    progress_callback: Callable[[int, str], None] | None = None,
) -> None:
    """Etapa 2: evaluaciones (entradas) + matriz bruta del cálculo."""
    doc.add_page_break()
    setattr(doc, '_report_table_page_used_cm', 0.0)

    alternativas_valores = _load_alternativas_valores(proyecto)
    if not alternativas_valores:
        _add_heading(doc, 'Etapa 2. Matrices de evaluación', level=1)
        _add_heading(doc, '2.1. Organización de la matriz', level=2)
        note = doc.add_paragraph()
        run = note.add_run('No hay alternativas definidas en el proyecto.')
        _set_run_font(run, color=RGBColor(0xB9, 0x1C, 0x1C))
        return

    evaluation_tables = [
        (dim, esc)
        for dim in (schema.get('dimensiones') or [])
        for esc in (dim.get('escenarios') or [])
        if _columnas_escenario_dimension(dim, esc.get('id'))
    ]
    total_evaluation_tables = max(len(evaluation_tables), 1)
    completed_evaluation_tables = 0

    section_21_pending = True
    section_22_pending = '2.2. Evaluaciones por dimensión y escenario'
    for dim in schema.get('dimensiones') or []:
        dim_nombre = dim.get('omoe_nombre') or f'Dimensión #{dim.get("omoe_id")}'
        escenarios = dim.get('escenarios') or []
        if not escenarios:
            continue
        dim_heading_pending = dim_nombre
        for esc in escenarios:
            esc_id = esc.get('id')
            esc_nombre = esc.get('nombre') or f'Escenario #{esc_id}'
            columnas = _columnas_escenario_dimension(dim, esc_id)
            if not columnas:
                continue
            headers = _evaluacion_matrix_headers(alternativas_valores)
            rows = _evaluacion_matrix_rows(alternativas_valores, columnas)
            # Contar títulos que se escriben ANTES de la tabla (tras el page-break
            # del empaquetado), para no dejar «Etapa 2» solo en una hoja.
            extra = 1  # título de escenario
            if section_21_pending:
                extra += 3  # Etapa 2 + 2.1 + párrafo
            if section_22_pending:
                extra += 1
            if dim_heading_pending:
                extra += 1
            _begin_table_block(
                doc,
                _estimate_evaluation_table_height_cm(
                    headers,
                    rows,
                    extra_heading_count=extra,
                ),
            )
            if section_21_pending:
                _add_heading(
                    doc,
                    'Etapa 2. Matrices de evaluación',
                    level=1,
                    keep_with_next=True,
                )
                _add_heading(
                    doc,
                    '2.1. Organización de la matriz',
                    level=2,
                    keep_with_next=True,
                )
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                run = p.add_run(
                    'Las tablas de evaluación se presentan transpuestas respecto al '
                    'módulo de carga: columnas = alternativas (eje x), filas = nodos '
                    'terminales (eje y). Los valores son las entradas x registradas '
                    'por el usuario. La matriz general (2.4) es el resultado agregado '
                    'del cálculo por dimensión, antes de Pareto, normalización o '
                    'método multicriterio.'
                )
                _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))
                section_21_pending = False
            if section_22_pending:
                _add_heading(doc, section_22_pending, level=2, keep_with_next=True)
                section_22_pending = None
            if dim_heading_pending:
                _add_heading(doc, dim_heading_pending, level=3, keep_with_next=True)
                dim_heading_pending = None
            _add_heading(doc, f'Escenario: {esc_nombre}', level=3, keep_with_next=True)
            _add_table(
                doc,
                headers,
                rows,
                title=f'Evaluación — {dim_nombre} — {esc_nombre}',
                new_page=False,
            )
            completed_evaluation_tables += 1
            if progress_callback:
                progress_callback(
                    84 + round(
                        5 * completed_evaluation_tables / total_evaluation_tables
                    ),
                    (
                        'Agregando tablas de evaluación '
                        f'({completed_evaluation_tables}/{total_evaluation_tables})'
                    ),
                )

    if section_21_pending:
        _add_heading(doc, 'Etapa 2. Matrices de evaluación', level=1)
        _add_heading(doc, '2.1. Organización de la matriz', level=2)
        note = doc.add_paragraph()
        run = note.add_run('No hay columnas de evaluación activas para mostrar.')
        _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))
    elif section_22_pending:
        _add_heading(doc, section_22_pending, level=2)
        note = doc.add_paragraph()
        run = note.add_run('No hay columnas de evaluación activas para mostrar.')
        _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))

    calculation_details: list[dict[str, Any]] = []
    bruta_headers, bruta_rows, bruta_error = _build_matriz_bruta_calculo(
        proyecto,
        progress_callback=progress_callback,
        details_out=calculation_details,
    )
    if bruta_error:
        _add_heading(doc, '2.4. Matriz general o matriz inicial', level=2)
        note = doc.add_paragraph()
        run = note.add_run(bruta_error)
        _set_run_font(run, color=RGBColor(0xB9, 0x1C, 0x1C))
        return

    _add_matriz_bruta_scientific_explanation(
        doc,
        calculation_details,
        progress_callback=progress_callback,
    )

    _begin_table_block(
        doc,
        _estimate_table_height_cm(bruta_headers, bruta_rows, extra_heading_count=2),
    )
    _add_heading(doc, '2.4. Matriz general o matriz inicial', level=2, keep_with_next=True)
    intro = doc.add_paragraph()
    intro.paragraph_format.keep_with_next = True
    run2 = intro.add_run(
        'Matriz que arroja el cálculo al inicio: alternativas × dimensiones con el valor '
        'agregado de cada dimensión, antes de filtro de Pareto, normalización o MADM '
        '(equivalente a la «matriz de utilidades» / matriz original del cálculo).'
    )
    _set_run_font(run2, color=RGBColor(0x47, 0x55, 0x69))
    _add_table(
        doc,
        bruta_headers,
        bruta_rows,
        title='Matriz bruta del cálculo (alternativas × dimensiones)',
        new_page=False,
    )


def build_informe_proyecto_docx(
    proyecto: Proyecto,
    *,
    include_map_weights: bool = False,
    progress_callback: Callable[[int, str], None] | None = None,
) -> bytes:
    """Informe de proyecto: Etapa 1 (estructura) + Etapa 2 (evaluaciones) en un solo documento."""
    def progress(percent: int, stage: str) -> None:
        if progress_callback:
            progress_callback(percent, stage)

    progress(3, 'Consultando estructura y evaluaciones')
    schema = build_evaluacion_schema(proyecto)
    progress(8, 'Preparando documento')
    doc = Document()
    _setup_document_fonts(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    _add_header_logos(doc)
    progress(12, 'Agregando información general')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Informe de proyecto')
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(proyecto.nombre)
    _set_run_font(run, bold=True)

    _add_heading(doc, 'Información general del proyecto')
    _add_table(
        doc,
        ['Campo', 'Valor'],
        [
            ['Nombre', proyecto.nombre],
            ['Descripción', proyecto.descripcion],
            ['Eslora máxima', proyecto.eslora_maxima],
            ['Desplazamiento', proyecto.desplazamiento],
            ['Velocidad máxima', proyecto.velocidad_maxima],
            ['Velocidad crucero', proyecto.velocidad_crucero],
            ['Tripulación', proyecto.tripulacion],
            ['Autonomía', proyecto.autonomia],
            ['Propulsión', proyecto.propulsion],
            ['Posicionamiento dinámico', proyecto.posicionamiento_dinamico],
            ['Laboratorios', proyecto.laboratorios],
            ['Otras características', proyecto.otras_caracteristicas],
        ],
        title='Información general del proyecto',
    )

    _add_etapa1_arboles_section(
        doc,
        schema,
        progress_callback=progress_callback,
    )
    progress(58, 'Agregando alternativas')
    _add_alternativas_section(doc, proyecto)
    progress(66, 'Agregando estructura jerárquica')
    _add_etapa1_estructura_jerarquica(doc, schema)
    progress(75, 'Agregando pesos de los nodos')
    _add_etapa1_pesos_nodo(doc, schema)
    progress(84, 'Construyendo matrices de evaluación')
    _add_etapa2_evaluaciones(
        doc,
        proyecto,
        schema,
        progress_callback=progress_callback,
    )
    progress(98, 'Guardando el documento Word')

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Documento generado por HATD · Informe de proyecto')
    _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    buf = BytesIO()
    doc.save(buf)
    progress(99, 'Preparando descarga')
    return buf.getvalue()


def build_informe_curvas_docx(proyecto: Proyecto) -> bytes:
    """Informe Word con curvas de utilidad finales (nodos terminales × escenario)."""
    from .curvas_utilidad_service import build_curvas_utilidad_export

    payload = build_curvas_utilidad_export(proyecto)
    doc = Document()
    _setup_document_fonts(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    _add_header_logos(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Curvas de utilidad finales')
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'Proyecto: {proyecto.nombre}')
    _set_run_font(run, bold=True)

    gen = payload.get('generado_en') or ''
    if gen:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f'Generado: {gen}')
        _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    nota = doc.add_paragraph()
    run = nota.add_run(
        'Inventario de funciones u(x) configuradas en nodos terminales por escenario. '
        'No es el historial de cambios: refleja el estado final del árbol para informes '
        'metodológicos. Dimensiones en modo valor bruto (p. ej. costos) suelen omitirse '
        'si no tienen familia de función.'
    )
    _set_run_font(run, color=RGBColor(0x47, 0x55, 0x69))

    dims = payload.get('dimensiones') or []
    total_curvas = sum(len(d.get('curvas') or []) for d in dims)
    sec_n = 2
    if total_curvas == 0:
        p = doc.add_paragraph()
        run = p.add_run(
            'No hay curvas de utilidad configuradas en este proyecto '
            '(sin familia de función en nodos terminales).'
        )
        _set_run_font(run, color=RGBColor(0xB9, 0x1C, 0x1C))
    else:
        doc.add_paragraph()
        h = doc.add_paragraph()
        run = h.add_run('1. Resumen')
        _set_run_font(run, bold=True)
        p = doc.add_paragraph()
        run = p.add_run(
            f'{total_curvas} curva(s) en {len(dims)} dimensión(es). '
            'Cada fila corresponde a un nodo terminal × escenario.'
        )
        _set_run_font(run)

        for dim in dims:
            curvas = dim.get('curvas') or []
            if not curvas:
                continue
            doc.add_paragraph()
            h = doc.add_paragraph()
            nombre = dim.get('omoe_nombre') or f"Dimensión #{dim.get('omoe_id')}"
            rama = (dim.get('rama_evaluacion') or '').upper() or '—'
            run = h.add_run(f'{sec_n}. {nombre} ({rama})')
            _set_run_font(run, bold=True)
            sec_n += 1

            meta_bits = []
            if dim.get('modo_valor_terminal'):
                meta_bits.append(f"modo valor={dim['modo_valor_terminal']}")
            if dim.get('escenario_agregacion'):
                meta_bits.append(f"agregación={dim['escenario_agregacion']}")
            if meta_bits:
                p = doc.add_paragraph()
                run = p.add_run(' · '.join(meta_bits))
                _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

            _add_table_caption(doc, f'Curvas de utilidad — {nombre} ({rama})')
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            _set_cell_text(hdr[0], 'Nodo terminal', bold=True)
            _set_cell_text(hdr[1], 'Escenario', bold=True)
            _set_cell_text(hdr[2], 'Familia', bold=True)
            _set_cell_text(hdr[3], 'Parámetros', bold=True)
            _set_cell_text(hdr[4], 'Unidad', bold=True)
            for c in curvas:
                cells = table.add_row().cells
                _set_cell_text(cells[0], str(c.get('terminal_nombre') or '—'))
                _set_cell_text(
                    cells[1],
                    str(c.get('escenario_nombre') or c.get('escenario_label') or '—'),
                )
                _set_cell_text(cells[2], str(c.get('familia_funciones') or '—'))
                _set_cell_text(
                    cells[3],
                    str(
                        c.get('constantes_display')
                        or _format_constantes(c.get('constantes'))
                        or '—'
                    ),
                )
                _set_cell_text(cells[4], str(c.get('unidad') or '—'))
            _finalize_report_table(table)

    doc.add_paragraph()
    h_ref = doc.add_paragraph()
    run = h_ref.add_run(f'{sec_n}. Nota metodológica')
    _set_run_font(run, bold=True)
    for text in (
        'La utilidad u(x) transforma el valor ofertado x en [0, 1] según la familia '
        'y constantes del nodo (capa micro).',
        'En simulación MADM, las utilidades de hojas se agregan con pesos del árbol; '
        'la resolución de escenarios (capa meso) usa Eq. (21)–(23) según la dimensión.',
        'Hay exportación JSON equivalente (curvas-utilidad) para procesamiento externo.',
    ):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        _set_run_font(run)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        'Documento generado por HATD · uso interno Cotecmar / ENAP / Universidad de la Costa'
    )
    _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_constantes(constantes: Any) -> str:
    if not constantes or not isinstance(constantes, dict):
        return ''
    parts = []
    for k, v in constantes.items():
        if v is None or v == '':
            continue
        parts.append(f'{k}={v}')
    return ', '.join(parts)
