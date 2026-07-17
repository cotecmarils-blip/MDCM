"""Trazas científicas y gráficos del Informe de resultados (Etapa 3)."""
from __future__ import annotations

import itertools
import math
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .informe_math import add_eq_line as _add_eq_line
from .informe_word_service import (
    _add_heading,
    _add_table,
    _begin_table_block,
    _set_run_font,
)

# Paleta pastel (estilo artículo científico) para todos los gráficos del informe.
_COLORS = [
    '#8FB8DE',  # azul pastel
    '#F4B183',  # naranja pastel
    '#A8D5A2',  # verde pastel
    '#E89AA6',  # rojo/rosa pastel
    '#C3B1E1',  # morado pastel
    '#F2D57E',  # amarillo pastel
    '#8FD0C4',  # turquesa pastel
    '#D9A6C2',  # magenta pastel
]

# Color primario pastel para barras de una sola serie.
_BAR_COLOR = '#9DBFE0'
_BAR_EDGE = '#5E82A6'

# Fondo amarillo pastel para gráficos de plano cartesiano (estilo paper).
_CARTESIAN_BG = '#FCF7E3'
_GRID_COLOR = '#E7DCB6'
_SPINE_COLOR = '#CDBE8E'


def _style_cartesian(ax, *, grid_axis: str = 'both') -> None:
    """Aplica fondo amarillo pastel y rejilla suave a ejes cartesianos."""
    ax.set_facecolor(_CARTESIAN_BG)
    for spine in ax.spines.values():
        spine.set_edgecolor(_SPINE_COLOR)
    if grid_axis == 'none':
        ax.grid(False)
    else:
        ax.grid(True, axis=grid_axis, alpha=0.55, color=_GRID_COLOR, linewidth=0.8)


def _pastel_cmap():
    """Colormap pastel secuencial (bajo→alto) para mapas de calor."""
    from matplotlib.colors import LinearSegmentedColormap

    return LinearSegmentedColormap.from_list(
        'pastel_seq',
        ['#FBEAA0', '#CFE3AE', '#9FD3C7', '#8FB8DE'],
    )


def _fmt_num(value: Any, digits: int = 4) -> str:
    if value is None or value == '':
        return '—'
    try:
        n = float(value)
        if math.isnan(n) or math.isinf(n):
            return '—'
        text = f'{n:.{digits}f}'.rstrip('0').rstrip('.')
        return text or '0'
    except (TypeError, ValueError):
        return str(value)


def _add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, color=RGBColor(0x33, 0x41, 0x55))
    run.italic = italic


def _is_min(direction: str) -> bool:
    return str(direction).strip().lower() in ('min', 'minimize', 'minimizar', 'menor')


def resolve_directions_for_dims(
    dims: list[dict[str, Any]],
    opciones: dict[str, Any],
) -> list[str]:
    direcciones = opciones.get('direcciones')
    dir_por_nombre = opciones.get('direcciones_por_dimension') or {}
    out: list[str] = []
    for i, d in enumerate(dims):
        dir_val = None
        if isinstance(dir_por_nombre, dict) and d['nombre'] in dir_por_nombre:
            dir_val = dir_por_nombre.get(d['nombre'])
        elif isinstance(direcciones, dict):
            dir_val = (
                direcciones.get(str(d['id']))
                or direcciones.get(d['id'])
                or direcciones.get(d['nombre'])
            )
        elif isinstance(direcciones, list) and i < len(direcciones):
            dir_val = direcciones[i]
        if dir_val is None:
            dir_val = 'min' if d.get('rama') in ('omoc', 'omor') else 'max'
        out.append('min' if _is_min(str(dir_val)) else 'max')
    return out


def extract_decision_matrix(
    resultado: dict[str, Any],
    dims: list[dict[str, Any]],
    *,
    only_pareto: bool = True,
) -> tuple[list[str], list[list[float]], list[int]]:
    """Retorna (nombres_alt, matriz, índices originales)."""
    alts = resultado.get('alternativas') or []
    idxs = []
    names = []
    rows = []
    for i, alt in enumerate(alts):
        if only_pareto and alt.get('excluida_pareto'):
            continue
        by_id = {
            d.get('omoe_id'): d.get('valor')
            for d in (alt.get('dimensiones') or [])
        }
        row = []
        ok = True
        for d in dims:
            v = by_id.get(d['id'])
            if v is None:
                ok = False
                break
            row.append(float(v))
        if not ok:
            continue
        idxs.append(i)
        nombre = (alt.get('nombre') or '').strip()
        apodo = (alt.get('apodo') or '').strip()
        if apodo and nombre:
            names.append(f'{nombre} ({apodo})')
        else:
            names.append(nombre or apodo or f'Alt {alt.get("id", i + 1)}')
        rows.append(row)

    # Prefer matriz_original filtrada si coincide en columnas
    raw = resultado.get('matriz_original')
    if raw and dims and len(raw) == len(alts) and raw and len(raw[0]) == len(dims):
        names2, rows2, idxs2 = [], [], []
        for i, alt in enumerate(alts):
            if only_pareto and alt.get('excluida_pareto'):
                continue
            if i >= len(raw):
                break
            nombre = (alt.get('nombre') or '').strip()
            apodo = (alt.get('apodo') or '').strip()
            if apodo and nombre:
                names2.append(f'{nombre} ({apodo})')
            else:
                names2.append(nombre or apodo or f'Alt {alt.get("id", i + 1)}')
            rows2.append([float(v) for v in raw[i]])
            idxs2.append(i)
        if rows2:
            return names2, rows2, idxs2

    return names, rows, idxs


def add_pareto_trace(
    doc: Document,
    *,
    matrix: list[list[float]],
    alt_names: list[str],
    dim_names: list[str],
    directions: list[str],
    epsilon: float = 0.0,
) -> None:
    """Desarrollo milimétrico del filtro Pareto a partir de la matriz de entrada."""
    if not matrix or not alt_names or not dim_names:
        return

    from .pareto_solver import ParetoSolver

    eps = float(epsilon) if epsilon is not None else 0.0
    _add_heading(doc, '3.3.1. Definición de dominancia', level=3)
    _add_paragraph(
        doc,
        'Una alternativa A domina a una alternativa B (con tolerancia ε) si A es '
        'al menos tan buena como B en todas las dimensiones y estrictamente mejor '
        'en al menos una. Para comparar dimensiones de maximizar y minimizar en '
        'la misma regla, primero se orienta la matriz al sentido «mayor es mejor».',
    )
    _add_eq_line(
        doc,
        'Orientación',
        r'\tilde{y}_{ij}=x_{ij}\ \mathrm{(max)},\quad'
        r'\tilde{y}_{ij}=-x_{ij}\ \mathrm{(min)}',
    )
    _add_eq_line(
        doc,
        'A domina a B',
        r'\forall j:\ \tilde{y}_{Aj}\geq\tilde{y}_{Bj}-\varepsilon'
        r'\quad\mathrm{y}\quad'
        r'\exists j:\ \tilde{y}_{Aj}>\tilde{y}_{Bj}+\varepsilon',
    )
    _add_eq_line(doc, 'ε usado en este cálculo', _fmt_num(eps, 12), latex=False)

    # --- Signos y matriz transformada ---
    _add_heading(doc, '3.3.2. Orientación de la matriz (sustitución)', level=3)
    sign_rows = []
    signs: list[float] = []
    for j, name in enumerate(dim_names):
        d = directions[j] if j < len(directions) else 'max'
        if _is_min(d):
            signs.append(-1.0)
            sign_rows.append([
                name,
                'min (menor es mejor)',
                'ỹ = −x',
                'signo = −1',
            ])
        else:
            signs.append(1.0)
            sign_rows.append([
                name,
                'max (mayor es mejor)',
                'ỹ = x',
                'signo = +1',
            ])
    _add_table(
        doc,
        ['Dimensión', 'Sentido', 'Transformación', 'Signo'],
        sign_rows,
        title='Orientación por dimensión',
        new_page=False,
    )

    # Matriz original (recordatorio)
    orig_rows = []
    for i, name in enumerate(alt_names):
        orig_rows.append([name, *[_fmt_num(v) for v in matrix[i]]])
    _add_table(
        doc,
        ['Alternativa', *dim_names],
        orig_rows,
        title='Matriz de entrada X (antes de Pareto)',
        new_page=False,
    )

    # Matriz transformada con sustitución por celda (compacta: una fila por alt×dim en tabla resumen)
    transformed: list[list[float]] = []
    subst_rows = []
    for i, name in enumerate(alt_names):
        trow = []
        for j, dim in enumerate(dim_names):
            x = float(matrix[i][j])
            s = signs[j]
            y = s * x
            trow.append(y)
            if s < 0:
                subst = f'ỹ = −({_fmt_num(x)}) = {_fmt_num(y)}'
            else:
                subst = f'ỹ = {_fmt_num(x)}'
            subst_rows.append([name, dim, _fmt_num(x), subst, _fmt_num(y)])
        transformed.append(trow)

    _add_table(
        doc,
        ['Alternativa', 'Dimensión', 'xᵢⱼ', 'Sustitución', 'ỹᵢⱼ'],
        subst_rows,
        title='Transformación celda a celda X → Ỹ',
        new_page=False,
    )
    y_rows = [[alt_names[i], *[_fmt_num(v) for v in transformed[i]]] for i in range(len(alt_names))]
    _add_table(
        doc,
        ['Alternativa', *dim_names],
        y_rows,
        title='Matriz orientada Ỹ (mayor es mejor en todas las columnas)',
        new_page=False,
    )

    # --- Ejecutar solver con trazas ---
    solver = ParetoSolver(
        matrix=matrix,
        dimensions=dim_names,
        directions=directions,
        alternatives=alt_names,
        epsilon=eps,
    )
    result = solver.solve(collect_all_dominators=True, trace=True)

    _add_heading(doc, '3.3.3. Por qué se excluye cada alternativa dominada', level=3)
    _add_paragraph(
        doc,
        'Para cada alternativa excluida se muestra al menos un dominador y la '
        'comparación dimensión a dimensión con los valores reales (en la escala '
        'orientada Ỹ y en la escala original X).',
    )

    # Index dominators: dominated -> list of dominators
    by_dominated: dict[str, list[str]] = {}
    for rel in result.dominance_relations:
        by_dominated.setdefault(rel.dominated_alternative, []).append(
            rel.dominator_alternative
        )

    name_to_idx = {n: i for i, n in enumerate(alt_names)}

    if not result.dominated_alternatives:
        _add_paragraph(
            doc,
            'Ninguna alternativa fue dominada: todas permanecen en el frente de Pareto.',
        )
    else:
        for dominated in result.dominated_alternatives:
            dominators = by_dominated.get(dominated) or []
            if not dominators:
                continue
            # Prefer the first dominator; show all names
            candidate = dominators[0]
            i_ref = name_to_idx.get(dominated)
            i_cand = name_to_idx.get(candidate)
            if i_ref is None or i_cand is None:
                continue

            _add_heading(
                doc,
                f'Exclusión de «{dominated}» (dominada por «{candidate}»'
                + (f' y otras: {", ".join(dominators[1:])}' if len(dominators) > 1 else '')
                + ')',
                level=3,
                keep_with_next=True,
            )
            _add_eq_line(
                doc,
                'Afirmación',
                f'«{candidate}» domina a «{dominated}» con ε = {_fmt_num(eps)}',
                latex=False,
            )

            dim_cmp_rows = []
            better_count = 0
            worse_count = 0
            equal_count = 0
            for j, dim in enumerate(dim_names):
                x_c = float(matrix[i_cand][j])
                x_r = float(matrix[i_ref][j])
                y_c = float(transformed[i_cand][j])
                y_r = float(transformed[i_ref][j])
                dirc = directions[j] if j < len(directions) else 'max'
                # En Ỹ: mayor es mejor
                if y_c > y_r + eps:
                    verdict = 'candidato mejor'
                    better_count += 1
                elif y_c < y_r - eps:
                    verdict = 'candidato peor'
                    worse_count += 1
                else:
                    verdict = 'equivalente (±ε)'
                    equal_count += 1
                dim_cmp_rows.append([
                    dim,
                    'min' if _is_min(dirc) else 'max',
                    _fmt_num(x_c),
                    _fmt_num(x_r),
                    _fmt_num(y_c),
                    _fmt_num(y_r),
                    f'{_fmt_num(y_c)} − {_fmt_num(y_r)} = {_fmt_num(y_c - y_r)}',
                    verdict,
                ])
            _add_table(
                doc,
                [
                    'Dimensión',
                    'Sentido',
                    f'x ({candidate})',
                    f'x ({dominated})',
                    f'ỹ ({candidate})',
                    f'ỹ ({dominated})',
                    'Δỹ = ỹ_c − ỹ_r',
                    'Veredicto',
                ],
                dim_cmp_rows,
                title=f'Comparación dimensional: {candidate} vs {dominated}',
                new_page=False,
            )
            _add_paragraph(
                doc,
                f'Reglas: (1) mejor o igual en todas → peor_count={worse_count} debe ser 0; '
                f'(2) estrictamente mejor en al menos una → better_count={better_count} ≥ 1. '
                f'Equivalencias (±ε): {equal_count}. '
                + (
                    'Se cumple → «'
                    + dominated
                    + '» queda excluida.'
                    if worse_count == 0 and better_count >= 1
                    else 'Condición de dominancia verificada por el solver.'
                ),
            )

    # Resumen de relaciones
    if result.dominance_relations:
        rel_rows = [
            [rel.dominator_alternative, rel.dominated_alternative]
            for rel in result.dominance_relations
        ]
        _add_table(
            doc,
            ['Domina (candidato)', 'Es dominada (excluida)'],
            rel_rows,
            title='Relaciones de dominancia encontradas',
            new_page=False,
        )

    _add_heading(doc, '3.3.4. Resultado del filtro', level=3)
    _add_table(
        doc,
        ['Estado', 'Alternativas', 'Criterio'],
        [
            [
                'No dominadas (continúan)',
                ', '.join(result.pareto_alternatives) if result.pareto_alternatives else '—',
                'Ninguna otra las domina',
            ],
            [
                'Dominadas (excluidas)',
                ', '.join(result.dominated_alternatives) if result.dominated_alternatives else '—',
                'Existe al menos un dominador',
            ],
        ],
        title='Frente de Pareto',
        new_page=False,
    )

    # Matriz de trabajo
    work_rows = []
    for i in result.pareto_indices:
        work_rows.append([alt_names[i], *[_fmt_num(v) for v in matrix[i]]])
    if work_rows:
        _add_table(
            doc,
            ['Alternativa', *dim_names],
            work_rows,
            title='Matriz de trabajo tras filtro Pareto (entra a normalización)',
            new_page=False,
        )


def add_normalization_trace(
    doc: Document,
    *,
    method: str,
    matrix: list[list[float]],
    alt_names: list[str],
    dim_names: list[str],
    directions: list[str],
) -> None:
    """Ecuación → sustitución con valores reales → resultado por celda."""
    if not matrix or not dim_names:
        return
    _add_heading(doc, '3.4.2. Desarrollo numérico (paso a paso)', level=3)
    _add_paragraph(
        doc,
        'A continuación se muestra cómo se obtiene cada columna de la matriz '
        'normalizada a partir de los valores de entrada: fórmula general, '
        'cálculo de parámetros de la columna y sustitución celda a celda.',
    )

    n = len(matrix)
    m = len(dim_names)
    for j, dim in enumerate(dim_names):
        if j >= m:
            break
        col = [row[j] for row in matrix if j < len(row)]
        if not col:
            continue
        direction = directions[j] if j < len(directions) else 'max'
        _add_heading(doc, f'Dimensión «{dim}» ({direction})', level=3, keep_with_next=True)

        if method == 'directional_minmax':
            vmin, vmax = min(col), max(col)
            rng = vmax - vmin
            if _is_min(direction):
                _add_eq_line(
                    doc,
                    'Fórmula',
                    r'r_{ij}=\dfrac{\max_j-x_{ij}}{\max_j-\min_j}',
                )
            else:
                _add_eq_line(
                    doc,
                    'Fórmula',
                    r'r_{ij}=\dfrac{x_{ij}-\min_j}{\max_j-\min_j}',
                )
            _add_eq_line(
                doc,
                'Parámetros de la columna',
                f'min = {_fmt_num(vmin)}, max = {_fmt_num(vmax)}, rango = {_fmt_num(rng)}',
                latex=False,
            )
            rows = []
            for i, x in enumerate(col):
                name = alt_names[i] if i < len(alt_names) else f'Alt {i + 1}'
                if abs(rng) < 1e-12:
                    r = 1.0
                    subst = f'rango ≈ 0 → r = 1'
                elif _is_min(direction):
                    r = (vmax - x) / rng
                    subst = (
                        f'r = ({_fmt_num(vmax)} − {_fmt_num(x)}) / {_fmt_num(rng)} '
                        f'= {_fmt_num(r, 6)}'
                    )
                else:
                    r = (x - vmin) / rng
                    subst = (
                        f'r = ({_fmt_num(x)} − {_fmt_num(vmin)}) / {_fmt_num(rng)} '
                        f'= {_fmt_num(r, 6)}'
                    )
                rows.append([name, _fmt_num(x), subst, _fmt_num(r, 6)])
            _add_table(
                doc,
                ['Alternativa', 'xᵢⱼ', 'Sustitución', 'rᵢⱼ'],
                rows,
                title=f'Desarrollo min-max — {dim}',
                new_page=False,
            )

        elif method in ('vector', 'directional_vector'):
            norm = math.sqrt(sum(x * x for x in col))
            _add_eq_line(
                doc,
                'Fórmula base',
                r'r_{ij}=\dfrac{x_{ij}}{\sqrt{\sum_i x_{ij}^{2}}}',
            )
            _add_eq_line(
                doc,
                'Norma de la columna',
                rf'\sqrt{{\sum x^{{2}}}}={_fmt_num(norm, 6)}',
            )
            rows = []
            for i, x in enumerate(col):
                name = alt_names[i] if i < len(alt_names) else f'Alt {i + 1}'
                base = (x / norm) if abs(norm) > 1e-12 else 0.0
                if method == 'directional_vector' and _is_min(direction):
                    r = 1.0 - base
                    subst = (
                        f'v = {_fmt_num(x)}/{_fmt_num(norm, 6)} = {_fmt_num(base, 6)}; '
                        f'r = 1 − {_fmt_num(base, 6)} = {_fmt_num(r, 6)}'
                    )
                else:
                    r = base
                    subst = f'r = {_fmt_num(x)} / {_fmt_num(norm, 6)} = {_fmt_num(r, 6)}'
                rows.append([name, _fmt_num(x), subst, _fmt_num(r, 6)])
            _add_table(
                doc,
                ['Alternativa', 'xᵢⱼ', 'Sustitución', 'rᵢⱼ'],
                rows,
                title=f'Desarrollo vectorial — {dim}',
                new_page=False,
            )

        elif method == 'sum':
            if _is_min(direction):
                inv = [1.0 / x if abs(x) > 1e-12 else 0.0 for x in col]
                total = sum(inv)
                _add_eq_line(
                    doc,
                    'Fórmula',
                    r'r_{ij}=\dfrac{1/x_{ij}}{\sum_k (1/x_{kj})}',
                )
                _add_eq_line(doc, 'Σ (1/x)', _fmt_num(total, 6), latex=False)
                rows = []
                for i, x in enumerate(col):
                    name = alt_names[i] if i < len(alt_names) else f'Alt {i + 1}'
                    r = (inv[i] / total) if abs(total) > 1e-12 else 0.0
                    subst = (
                        f'r = (1/{_fmt_num(x)}) / {_fmt_num(total, 6)} = {_fmt_num(r, 6)}'
                    )
                    rows.append([name, _fmt_num(x), subst, _fmt_num(r, 6)])
            else:
                total = sum(col)
                _add_eq_line(
                    doc,
                    'Fórmula',
                    r'r_{ij}=\dfrac{x_{ij}}{\sum_i x_{ij}}',
                )
                _add_eq_line(doc, 'Σ x', _fmt_num(total, 6), latex=False)
                rows = []
                for i, x in enumerate(col):
                    name = alt_names[i] if i < len(alt_names) else f'Alt {i + 1}'
                    r = (x / total) if abs(total) > 1e-12 else 0.0
                    subst = f'r = {_fmt_num(x)} / {_fmt_num(total, 6)} = {_fmt_num(r, 6)}'
                    rows.append([name, _fmt_num(x), subst, _fmt_num(r, 6)])
            _add_table(
                doc,
                ['Alternativa', 'xᵢⱼ', 'Sustitución', 'rᵢⱼ'],
                rows,
                title=f'Desarrollo por suma — {dim}',
                new_page=False,
            )
        else:
            _add_paragraph(
                doc,
                f'Método «{method}»: se muestra la matriz normalizada resultante; '
                'el desarrollo simbólico detallado se limita a los métodos documentados.',
                italic=True,
            )
            break

        # Limitar verbosidad: máximo 4 dimensiones desarrolladas + resumen
        if j >= 3 and m > 4:
            _add_paragraph(
                doc,
                f'Se desarrollaron las primeras 4 dimensiones; las restantes ({m - 4}) '
                'siguen el mismo procedimiento (véase matriz normalizada).',
                italic=True,
            )
            break


def add_weights_trace(
    doc: Document,
    *,
    method: str,
    weights: list[float],
    dim_names: list[str],
    norm_matrix: list[list[float]] | None = None,
    pesos_usuario: list[float] | None = None,
) -> None:
    _add_heading(doc, '3.5.2. Desarrollo numérico de los pesos', level=3)
    m = len(dim_names)
    if method == 'equal_weights':
        _add_eq_line(doc, 'Fórmula', r'w_{j}=\dfrac{1}{m}')
        _add_eq_line(
            doc,
            'Sustitución',
            f'm = {m} → wⱼ = 1/{m} = {_fmt_num(1.0 / m if m else 0, 6)}',
            latex=False,
        )
        rows = [[dim_names[j] if j < len(dim_names) else f'D{j+1}', _fmt_num(weights[j] if j < len(weights) else 1/m, 6)]
                for j in range(m)]
        _add_table(doc, ['Dimensión', 'wⱼ'], rows, title='Pesos iguales — sustitución', new_page=False)
        return

    if method == 'user_defined_weights':
        _add_eq_line(doc, 'Fórmula', r'w_{j}=\dfrac{p_{j}}{\sum_k p_{k}}')
        raw = pesos_usuario or [w * 100 for w in weights]
        total = sum(float(p) for p in raw)
        _add_eq_line(doc, 'Σ pⱼ', f'{_fmt_num(total, 4)} %', latex=False)
        rows = []
        for j, name in enumerate(dim_names):
            p = float(raw[j]) if j < len(raw) else 0.0
            w = (p / total) if abs(total) > 1e-12 else 0.0
            rows.append([
                name,
                f'{_fmt_num(p, 2)} %',
                f'w = {_fmt_num(p, 2)} / {_fmt_num(total, 2)} = {_fmt_num(w, 6)}',
                _fmt_num(weights[j] if j < len(weights) else w, 6),
            ])
        _add_table(
            doc,
            ['Dimensión', 'pⱼ', 'Sustitución', 'wⱼ'],
            rows,
            title='Pesos usuario — sustitución',
            new_page=False,
        )
        return

    if method == 'entropy' and norm_matrix:
        n = len(norm_matrix)
        _add_eq_line(doc, 'Proporción', r'p_{ij}=\dfrac{r_{ij}}{\sum_i r_{ij}}')
        _add_eq_line(
            doc,
            'Entropía',
            r'e_{j}=-\dfrac{\sum_i p_{ij}\,\ln(p_{ij})}{\ln(n)}',
        )
        _add_eq_line(
            doc,
            'Peso',
            r'w_{j}=\dfrac{1-e_{j}}{\sum_k (1-e_{k})}',
        )
        _add_eq_line(doc, 'n (alternativas)', str(n), latex=False)
        rows = []
        ones_minus_e = []
        e_list = []
        for j, name in enumerate(dim_names):
            col = [row[j] for row in norm_matrix if j < len(row)]
            s = sum(col)
            probs = [(v / s) if abs(s) > 1e-12 else 0.0 for v in col]
            ent = 0.0
            for p in probs:
                if p > 1e-12:
                    ent -= p * math.log(p)
            if n > 1:
                ent /= math.log(n)
            e_list.append(ent)
            ones_minus_e.append(1.0 - ent)
        denom = sum(ones_minus_e)
        for j, name in enumerate(dim_names):
            w = (ones_minus_e[j] / denom) if abs(denom) > 1e-12 else 0.0
            rows.append([
                name,
                _fmt_num(e_list[j], 6),
                _fmt_num(ones_minus_e[j], 6),
                f'w = {_fmt_num(ones_minus_e[j], 6)} / {_fmt_num(denom, 6)} = {_fmt_num(w, 6)}',
                _fmt_num(weights[j] if j < len(weights) else w, 6),
            ])
        _add_table(
            doc,
            ['Dimensión', 'eⱼ', '1−eⱼ', 'Sustitución', 'wⱼ'],
            rows,
            title='Entropía — desarrollo',
            new_page=False,
        )
        return

    # Fallback genérico
    _add_paragraph(
        doc,
        f'Método «{method}»: se reportan los pesos resultantes del cálculo.',
    )
    rows = [
        [dim_names[j] if j < len(dim_names) else f'D{j+1}', _fmt_num(weights[j], 6)]
        for j in range(len(weights))
    ]
    if rows:
        _add_table(doc, ['Dimensión', 'wⱼ'], rows, title='Pesos resultantes', new_page=False)


def compute_topsis_trace(
    norm_matrix: list[list[float]],
    weights: list[float],
) -> dict[str, Any]:
    """Traza TOPSIS asumiendo matriz ya orientada a beneficio."""
    n = len(norm_matrix)
    m = len(norm_matrix[0]) if n else 0
    V = []
    for i in range(n):
        V.append([
            float(norm_matrix[i][j]) * float(weights[j] if j < len(weights) else 0.0)
            for j in range(m)
        ])
    ideal_pos = [max(V[i][j] for i in range(n)) for j in range(m)]
    ideal_neg = [min(V[i][j] for i in range(n)) for j in range(m)]
    d_pos, d_neg, closeness = [], [], []
    for i in range(n):
        dp = math.sqrt(sum((V[i][j] - ideal_pos[j]) ** 2 for j in range(m)))
        dn = math.sqrt(sum((V[i][j] - ideal_neg[j]) ** 2 for j in range(m)))
        d_pos.append(dp)
        d_neg.append(dn)
        closeness.append(dn / (dp + dn) if (dp + dn) > 1e-12 else 0.0)
    return {
        'V': V,
        'ideal_pos': ideal_pos,
        'ideal_neg': ideal_neg,
        'd_pos': d_pos,
        'd_neg': d_neg,
        'closeness': closeness,
    }


def add_madm_trace(
    doc: Document,
    *,
    method: str,
    norm_matrix: list[list[float]],
    weights: list[float],
    alt_names: list[str],
    dim_names: list[str],
) -> None:
    _add_heading(doc, '3.6.2. Desarrollo numérico del ranking', level=3)
    if not norm_matrix or not weights:
        _add_paragraph(doc, 'No hay matriz normalizada/pesos suficientes para el desarrollo.', italic=True)
        return

    if method == 'topsis':
        _add_paragraph(
            doc,
            'Se construye la matriz ponderada Vᵢⱼ = wⱼ · rᵢⱼ; luego las soluciones '
            'ideales, las distancias euclídeas y la cercanía relativa Cᵢ.',
        )
        trace = compute_topsis_trace(norm_matrix, weights)
        # Mostrar V para primera alternativa como ejemplo + tabla completa compacta
        _add_eq_line(doc, 'Matriz ponderada', r'v_{ij}=w_{j}\,r_{ij}')
        v_rows = []
        for i, name in enumerate(alt_names):
            if i >= len(trace['V']):
                break
            cells = [_fmt_num(v, 6) for v in trace['V'][i]]
            # Sustitución corta de la primera celda
            if dim_names and weights and norm_matrix[i]:
                w0 = weights[0]
                r0 = norm_matrix[i][0]
                subst0 = f'v₁ = {_fmt_num(w0, 6)}·{_fmt_num(r0, 6)} = {_fmt_num(w0 * r0, 6)}'
            else:
                subst0 = '—'
            v_rows.append([name, subst0, *cells])
        _add_table(
            doc,
            ['Alternativa', 'Ej. 1ª col.', *dim_names],
            v_rows,
            title='Matriz ponderada V (TOPSIS)',
            new_page=False,
        )
        _add_table(
            doc,
            ['Tipo', *dim_names],
            [
                ['Ideal + (A⁺)', *[_fmt_num(v, 6) for v in trace['ideal_pos']]],
                ['Ideal − (A⁻)', *[_fmt_num(v, 6) for v in trace['ideal_neg']]],
            ],
            title='Soluciones ideales',
            new_page=False,
        )
        _add_eq_line(
            doc,
            'Distancias',
            r'D_{i}^{+}=\sqrt{\sum_j(v_{ij}-v_{j}^{+})^{2}}\ ;\quad'
            r'D_{i}^{-}=\sqrt{\sum_j(v_{ij}-v_{j}^{-})^{2}}',
        )
        _add_eq_line(
            doc,
            'Cercanía',
            r'C_{i}=\dfrac{D_{i}^{-}}{D_{i}^{+}+D_{i}^{-}}',
        )
        dist_rows = []
        for i, name in enumerate(alt_names):
            if i >= len(trace['closeness']):
                break
            dp, dn, c = trace['d_pos'][i], trace['d_neg'][i], trace['closeness'][i]
            subst = (
                f'C = {_fmt_num(dn, 6)} / ({_fmt_num(dp, 6)} + {_fmt_num(dn, 6)}) '
                f'= {_fmt_num(c, 6)}'
            )
            dist_rows.append([
                name,
                _fmt_num(dp, 6),
                _fmt_num(dn, 6),
                subst,
                _fmt_num(c, 6),
            ])
        _add_table(
            doc,
            ['Alternativa', 'D⁺', 'D⁻', 'Sustitución Cᵢ', 'Cᵢ'],
            dist_rows,
            title='Distancias y cercanía relativa TOPSIS',
            new_page=False,
        )
        return

    if method == 'wsm':
        _add_eq_line(doc, 'Fórmula', r'S_{i}=\sum_j w_{j}\,r_{ij}')
        rows = []
        for i, name in enumerate(alt_names):
            if i >= len(norm_matrix):
                break
            terms = []
            total = 0.0
            for j, dim in enumerate(dim_names):
                w = float(weights[j]) if j < len(weights) else 0.0
                r = float(norm_matrix[i][j]) if j < len(norm_matrix[i]) else 0.0
                term = w * r
                total += term
                terms.append(f'{_fmt_num(w, 4)}·{_fmt_num(r, 4)}')
            subst = ' + '.join(terms) + f' = {_fmt_num(total, 6)}'
            rows.append([name, subst, _fmt_num(total, 6)])
        _add_table(
            doc,
            ['Alternativa', 'Sustitución Sᵢ', 'Sᵢ'],
            rows,
            title='WSM — desarrollo',
            new_page=False,
        )
        return

    _add_paragraph(
        doc,
        f'Método «{method}»: se reporta el ranking final. El desarrollo simbólico '
        'detallado con sustitución está disponible para TOPSIS y WSM.',
        italic=True,
    )


def _plt():
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    return plt


def _png_size_px(png: BytesIO) -> tuple[int, int] | None:
    """(ancho, alto) en píxeles leyendo la cabecera IHDR del PNG, sin PIL."""
    try:
        data = png.getvalue()
        if len(data) < 24 or data[:8] != b'\x89PNG\r\n\x1a\n':
            return None
        width = int.from_bytes(data[16:20], 'big')
        height = int.from_bytes(data[20:24], 'big')
        if width <= 0 or height <= 0:
            return None
        return width, height
    except Exception:
        return None


def add_picture_block(
    doc: Document,
    png: BytesIO,
    caption: str,
    *,
    height_cm: float | None = None,
    width_cm: float = 12.0,
) -> None:
    """Inserta figura y estima su altura real para empaquetar varias por página.

    La altura de empaquetado se calcula a partir de la relación de aspecto real
    del PNG (ancho renderizado × alto/ancho) más el espacio del pie de figura.
    Así las figuras «anchas y bajas» (barras horizontales, líneas) permiten 3+
    por hoja, mientras que las altas (3D, radar) siguen ocupando lo necesario.
    """
    if height_cm is None:
        size = _png_size_px(png)
        if size:
            w_px, h_px = size
            # Alto de la imagen renderizada + pie de figura (~0.7 cm).
            height_cm = width_cm * (h_px / w_px) + 0.7
        else:
            height_cm = 8.0
    # Umbrales de empaquetado propios de figuras: residuo mínimo pequeño para
    # permitir 3+ gráficos por hoja cuando la figura es baja.
    _begin_table_block(
        doc,
        estimated_height_cm=height_cm,
        min_remaining_cm=3.5,
        buffer_cm=0.4,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(png, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(caption)
    _set_run_font(r, size_pt=8)
    r.italic = True


def bar_1d_png(
    puntos: list[dict[str, Any]],
    dim_id: Any,
    dim_label: str,
    *,
    horizontal: bool = True,
) -> BytesIO | None:
    try:
        plt = _plt()
    except ImportError:
        return None
    names, vals = [], []
    for p in puntos:
        v = (p.get('valores') or {}).get(dim_id)
        if v is None:
            continue
        names.append(p.get('label') or p.get('nombre') or '')
        vals.append(float(v))
    if len(vals) < 1:
        return None
    fig, ax = plt.subplots(figsize=(5.4, max(2.2, min(3.6, 0.32 * len(names) + 0.8))), dpi=120)
    if horizontal:
        ax.barh(
            names[::-1], vals[::-1],
            color=_BAR_COLOR, edgecolor=_BAR_EDGE, linewidth=0.8,
        )
        ax.set_xlabel(dim_label)
        _style_cartesian(ax, grid_axis='x')
    else:
        ax.bar(
            names, vals,
            color=_BAR_COLOR, edgecolor=_BAR_EDGE, linewidth=0.8,
        )
        ax.set_ylabel(dim_label)
        ax.tick_params(axis='x', rotation=30)
        _style_cartesian(ax, grid_axis='y')
    ax.set_title(f'1D — {dim_label}', fontsize=11)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def scatter_2d_png(puntos, x_key, y_key, x_label, y_label, title) -> BytesIO | None:
    try:
        plt = _plt()
    except ImportError:
        return None
    xs, ys, labels = [], [], []
    for p in puntos:
        xv = (p.get('valores') or {}).get(x_key)
        yv = (p.get('valores') or {}).get(y_key)
        if xv is None or yv is None:
            continue
        xs.append(float(xv))
        ys.append(float(yv))
        labels.append(p.get('label') or p.get('nombre') or '')
    if len(xs) < 2:
        return None
    fig, ax = plt.subplots(figsize=(5.4, 3.6), dpi=120)
    _style_cartesian(ax)
    for i, (x, y, lab) in enumerate(zip(xs, ys, labels)):
        ax.scatter(
            x, y, s=70, color=_COLORS[i % len(_COLORS)],
            edgecolors='#5B6472', linewidths=0.7, zorder=3,
        )
        ax.annotate(lab, (x, y), textcoords='offset points', xytext=(5, 5), fontsize=8)
    ax.set_xlabel(x_label)
    ax.set_ylabel(y_label)
    ax.set_title(title, fontsize=11)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def scatter_2d_heat_png(
    puntos,
    x_key,
    y_key,
    color_key,
    x_label,
    y_label,
    color_label,
    title,
) -> BytesIO | None:
    """Espacio de decisión 2D con mapa de calor.

    Ejes X/Y son dos dimensiones y el color (barra) corresponde a una tercera
    dimensión (`color_key`), no a la puntuación MADM.
    """
    try:
        plt = _plt()
    except ImportError:
        return None
    xs, ys, cs, labels = [], [], [], []
    for p in puntos:
        vals = p.get('valores') or {}
        xv = vals.get(x_key)
        yv = vals.get(y_key)
        cv = vals.get(color_key)
        if xv is None or yv is None or cv is None:
            continue
        xs.append(float(xv))
        ys.append(float(yv))
        cs.append(float(cv))
        labels.append(p.get('label') or p.get('nombre') or '')
    if len(xs) < 2:
        return None
    fig, ax = plt.subplots(figsize=(5.6, 3.8), dpi=120)
    _style_cartesian(ax)
    sc = ax.scatter(
        xs, ys, c=cs, cmap=_pastel_cmap(), s=110,
        edgecolors='#5B6472', linewidths=0.8, zorder=3,
    )
    for x, y, lab in zip(xs, ys, labels):
        ax.annotate(lab, (x, y), textcoords='offset points', xytext=(5, 5), fontsize=8)
    cbar = fig.colorbar(sc, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label(color_label, fontsize=9)
    cbar.outline.set_edgecolor(_SPINE_COLOR)
    ax.set_xlabel(x_label)
    ax.set_ylabel(y_label)
    ax.set_title(title, fontsize=11)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def scatter_3d_png(puntos, x_key, y_key, z_key, x_label, y_label, z_label, title) -> BytesIO | None:
    try:
        plt = _plt()
        from mpl_toolkits.mplot3d import Axes3D  # noqa: F401
    except ImportError:
        return None
    xs, ys, zs, labels = [], [], [], []
    for p in puntos:
        xv = (p.get('valores') or {}).get(x_key)
        yv = (p.get('valores') or {}).get(y_key)
        zv = (p.get('valores') or {}).get(z_key)
        if xv is None or yv is None or zv is None:
            continue
        xs.append(float(xv))
        ys.append(float(yv))
        zs.append(float(zv))
        labels.append(p.get('label') or p.get('nombre') or '')
    if len(xs) < 2:
        return None
    fig = plt.figure(figsize=(5.6, 4.2), dpi=120)
    ax = fig.add_subplot(111, projection='3d')
    pane_rgba = (0.988, 0.969, 0.890, 0.55)  # amarillo pastel translúcido
    for axis in (ax.xaxis, ax.yaxis, ax.zaxis):
        axis.set_pane_color(pane_rgba)
        axis.pane.set_edgecolor(_SPINE_COLOR)
    for i, (x, y, z, lab) in enumerate(zip(xs, ys, zs, labels)):
        ax.scatter(
            x, y, z, s=55, color=_COLORS[i % len(_COLORS)],
            edgecolors='#5B6472', linewidths=0.6,
        )
        ax.text(x, y, z, f' {lab}', fontsize=7)
    ax.set_xlabel(x_label[:28])
    ax.set_ylabel(y_label[:28])
    ax.set_zlabel(z_label[:28])
    ax.set_title(title, fontsize=10)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def radar_png(puntos: list[dict[str, Any]], dims: list[dict[str, Any]]) -> BytesIO | None:
    if len(dims) < 2 or len(puntos) < 1:
        return None
    try:
        plt = _plt()
        import numpy as np
    except ImportError:
        return None
    labels = [d['nombre'] for d in dims]
    angles = np.linspace(0, 2 * np.pi, len(dims), endpoint=False).tolist()
    angles += angles[:1]

    # Recolectar las series con valores completos por dimensión.
    series = []
    for p in puntos[:8]:
        raw = []
        ok = True
        for d in dims:
            v = (p.get('valores') or {}).get(d['id'])
            if v is None:
                ok = False
                break
            raw.append(float(v))
        if not ok:
            continue
        series.append((p.get('label') or p.get('nombre'), raw))
    if not series:
        return None

    # Normalización min–max por eje: cada dimensión suele tener escalas y
    # unidades distintas (p. ej. Costo ~250 vs Efectividad ~0-1). Sin esto,
    # el eje de mayor magnitud domina y el resto se aplasta hacia el centro,
    # dando la falsa impresión de "un solo eje".
    n_dims = len(dims)
    col_min = [min(s[1][j] for s in series) for j in range(n_dims)]
    col_max = [max(s[1][j] for s in series) for j in range(n_dims)]

    def _norm(v: float, j: int) -> float:
        lo, hi = col_min[j], col_max[j]
        if hi <= lo:
            return 0.5  # sin variación en el eje: punto medio
        return (v - lo) / (hi - lo)

    fig, ax = plt.subplots(figsize=(5.4, 4.4), dpi=120, subplot_kw=dict(polar=True))
    for i, (label, raw) in enumerate(series):
        vals = [_norm(raw[j], j) for j in range(n_dims)]
        vals += vals[:1]
        color = _COLORS[i % len(_COLORS)]
        ax.plot(angles, vals, color=color, linewidth=1.6, label=label)
        ax.fill(angles, vals, color=color, alpha=0.08)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([lb[:18] for lb in labels], fontsize=8)
    ax.set_ylim(0, 1)
    ax.set_yticks([0.25, 0.5, 0.75, 1.0])
    ax.set_yticklabels(['25%', '50%', '75%', '100%'], fontsize=7, color=_SPINE_COLOR)
    ax.set_title('Radar — dimensiones del cálculo (normalizado por eje)', fontsize=11)
    ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.1), fontsize=7)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def ranking_bar_png(ranking_rows: list[tuple[str, float]]) -> BytesIO | None:
    if not ranking_rows:
        return None
    try:
        plt = _plt()
    except ImportError:
        return None
    names = [r[0] for r in ranking_rows][::-1]
    scores = [r[1] for r in ranking_rows][::-1]
    fig, ax = plt.subplots(figsize=(5.6, max(2.4, min(3.8, 0.36 * len(names) + 1.0))), dpi=120)
    _style_cartesian(ax, grid_axis='x')
    ax.barh(names, scores, color=_BAR_COLOR, edgecolor=_BAR_EDGE, linewidth=0.8)
    ax.set_xlabel('Puntuación MADM')
    ax.set_title('Ranking del cálculo', fontsize=11)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def ranking_util_madm_png(
    puntos: list[dict[str, Any]],
    *,
    madm_label: str = 'MADM',
) -> BytesIO | None:
    """Línea por ranking: Utilidad (valor global) vs puntuación MADM (TOPSIS…).

    Las alternativas se ordenan por ranking y se comparan dos series por
    alternativa: la utilidad agregada del árbol y el score del método MADM.
    """
    try:
        plt = _plt()
    except ImportError:
        return None

    filas = []
    for p in puntos:
        util = p.get('valor_global')
        score = p.get('score_madm')
        if util is None and score is None:
            continue
        filas.append({
            'label': p.get('label') or p.get('nombre') or '',
            'ranking': p.get('ranking'),
            'util': float(util) if util is not None else None,
            'score': float(score) if score is not None else None,
        })
    if len(filas) < 2:
        return None

    filas.sort(key=lambda f: (f['ranking'] is None, f['ranking'] if f['ranking'] is not None else 9999))

    x = list(range(len(filas)))
    labels = [f['label'] for f in filas]
    util_vals = [f['util'] for f in filas]
    score_vals = [f['score'] for f in filas]
    tiene_util = any(v is not None for v in util_vals)
    tiene_score = any(v is not None for v in score_vals)
    if not (tiene_util and tiene_score):
        return None

    fig, ax = plt.subplots(figsize=(5.8, 3.8), dpi=120)
    _style_cartesian(ax)
    ax.plot(
        x, util_vals, marker='o', markersize=6, linewidth=2.0,
        color=_COLORS[0], markeredgecolor='#5B6472', markeredgewidth=0.6,
        label='Utilidad (valor global)',
    )
    ax.plot(
        x, score_vals, marker='s', markersize=6, linewidth=2.0,
        color=_COLORS[3], markeredgecolor='#5B6472', markeredgewidth=0.6,
        label=f'Puntuación {madm_label}',
    )
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_xlabel('Alternativas (orden por ranking →)')
    ax.set_ylabel('Valor')
    ax.set_title(f'Ranking: Utilidad vs {madm_label}', fontsize=11)
    ax.legend(fontsize=8, loc='best', framealpha=0.85)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='PNG', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf
